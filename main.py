from PyQt6.QtWidgets import QApplication , QWidget,QVBoxLayout,QPushButton,QCheckBox,QLineEdit,QLabel,QMessageBox, \
                                                        QTextEdit,QFrame,QTableWidget,QGridLayout,QTableWidgetItem,QTreeView,\
                                                        QRadioButton,QComboBox,QFileDialog,QScrollArea,QMainWindow,QDateEdit,QTextEdit
from PyQt6.QtGui import QIcon,QFont,QIntValidator,QScreen,QFont,QPixmap,QColor,QMovie
from PyQt6 import uic
from PyQt6.QtCore import Qt,QCoreApplication,QDate,QLocale,QTimer,QSize
import sqlite3
import sys
from email.message import EmailMessage
import ssl
import smtplib
from datetime import date
from hijri_converter import Hijri, Gregorian
import os
import docx
import docx2pdf
from docx.shared import Mm
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.ns import qn as qn2
import shutil
from PIL import Image
import ctypes
from math import *
import comtypes.client
import aspose.words as aw


n = 0
def is_admin():
    try:
        return ctypes.windll.shell32.IsUserAnAdmin()
    except:
        return False
def mainToRunApp():
    global con,cr,app,window,n,sop1
    n +=1
    
    con = sqlite3.connect("app.db")
    cr = con.cursor()
    if n <=1:
        app = QApplication(sys.argv)
    window = Login()
    sop1 = scrollWindow()
    window.show()
    if n<=1:
        app.exec()

if is_admin():
    
    CURRENT_TIME_GEO = (str(date.today())).replace("-","/")
    CURRENT_TIME_HIJ_TEMP = (str(date.today())).split("-")
    CURRENT_TIME_HIJ = str(Gregorian(int(CURRENT_TIME_HIJ_TEMP[0]),int(CURRENT_TIME_HIJ_TEMP[1]),int(CURRENT_TIME_HIJ_TEMP[2])).to_hijri()).replace("-","/")
    title ="برنامج ادارة متطوعين"
    icon = "assests/icon.ico"
    
    class scrollWindow(QMainWindow):
        def __init__(self):
            super().__init__()
            self.resize(1120,680)
            self.setWindowTitle(title)
            self.setWindowIcon(QIcon(icon))
            self.setStyleSheet("background-color:white")
    class tempwindows(QWidget):
        def __init__(self):
            super().__init__()
            self.setWindowTitle(title)
            self.setWindowIcon(QIcon(icon))
    class Login(QWidget):
        def __init__(self):
            super().__init__()
            self.ob = 0
            self.changePasswordWindow = 0
            self.changeEmailWindow = 0
            uic.loadUi("login.ui",self)
            self.setWindowTitle(title)
            self.setWindowIcon(QIcon(icon))
            self.login.clicked.connect(self.loginPasswordFun)
            self.login.setShortcut("Return")
            self.backUp.clicked.connect(self.sendPasswordFun)
            self.passEntry.setEchoMode(QLineEdit.EchoMode.Password)
            self.changePass.clicked.connect(self.changePasswordFun)
            self.changeEmail.clicked.connect(self.changeEmailFun)
            self.developer.clicked.connect(self.developerFun)
            self.setFixedSize(676,452)
            
        def loginPasswordFun(self):
            cr.execute("SELECT password from loginStuff")
            password = cr.fetchone()[0]
            if password == self.passEntry.text():
                self.destroy()
                self.ob = Window()
            else:
                msg = QMessageBox(parent=self,text="كلمة المرور خاطئة")
                msg.setWindowTitle("ERROR")
                msg.setStyleSheet("background-color:white")
                msg.setIcon(QMessageBox.Icon.Critical)
                msg.exec()
        def sendPasswordFun(self):
            cr.execute("SELECT email from loginStuff")
            email_sender = "abodi3313@gmail.com"
            email_pass = "dduviygcxsmyxekr"

            email_recevier = f"{cr.fetchone()[0]}"
            subjet = "كلمة المرور"
            cr.execute("SELECT password from loginStuff")
            body = f"""
            اهلا وسهلا بك في البرنامج 
            كلمة سر الدخول هي{cr.fetchone()[0]}
            """

            em = EmailMessage()
            em['From']=email_sender
            em['To']=email_recevier
            em['subject'] = subjet
            em.set_content(body)

            context = ssl.create_default_context()

            with smtplib.SMTP_SSL('smtp.gmail.com',465,context=context) as smtp:
                smtp.login(email_sender, email_pass)
                smtp.sendmail(email_sender, email_recevier, em.as_string())
            d = QMessageBox(parent=self,text="تم الارسال بنجاح")
            d.setWindowTitle("نجاح")
            d.setIcon(QMessageBox.Icon.Information)
            d.setStyleSheet("background-color:white")
            ret = d.exec()
        def changeEmailFun(self):
            self.changeEmailWindow = tempwindows()
            self.changeEmailWindow.setWindowTitle(title)
            self.changeEmailWindow.setWindowIcon(QIcon(icon))
            self.changeEmailWindow.setFixedSize(280,250)

            passwordLabel = QLabel("كلمة المرور",self.changeEmailWindow)
            passwordLabel.setStyleSheet("font-size:25px")
            passwordLabel.move(70,10)

            self.PasswordToEmail = QLineEdit(self.changeEmailWindow)
            self.PasswordToEmail.setStyleSheet("font-size:20px;width:240px")
            self.PasswordToEmail.move(10,50)
            self.PasswordToEmail.setEchoMode(QLineEdit.EchoMode.Password)

            newEmail_label1 = QLabel("الايميل الجديد",self.changeEmailWindow)
            newEmail_label1.setStyleSheet("font-size:25px")
            newEmail_label1.move(70,90)
            self.newEmail = QLineEdit(self.changeEmailWindow)
            self.newEmail.setStyleSheet("font-size:20px;width:240px")
            self.newEmail.move(10,130)

            submitButton = QPushButton("تغيير",self.changeEmailWindow,clicked=self.changeEmailComplete)
            submitButton.setStyleSheet("font-size:20px;width:150px;background-color:green")
            submitButton.move(60,180)

            self.changeEmailWindow.show()
        def changeEmailComplete(self):
            cr.execute("SELECT password FROM loginStuff")
            curentPassword = cr.fetchone()[0]
            if self.PasswordToEmail.text() == curentPassword:
                if len(self.newEmail.text()) != 0:
                    cr.execute(f"UPDATE loginStuff SET email = '{self.newEmail.text()}'")
                    d = QMessageBox(parent =self.changeEmailWindow , text="تم تغيير كلمة المرور بنجاح")
                    d.setWindowTitle("نجاح")
                    d.setIcon(QMessageBox.Icon.Information)
                    d.setStyleSheet("background-color:white")
                    ret = d.exec()
                    con.commit()
                    self.changeEmailWindow.destroy()
            else:
                d = QMessageBox(parent=self.changeEmailWindow,text="كلمة المرور خاطئة")
                d.setWindowTitle("ERROR")
                d.setIcon(QMessageBox.Icon.Critical)
                d.setStyleSheet("background-color:white")
                ret = d.exec()
        def changePasswordFun(self):
            self.changePasswordWindow = tempwindows()
            self.changePasswordWindow.setWindowTitle(title)
            self.changePasswordWindow.setWindowIcon(QIcon(icon))
            self.changePasswordWindow.setFixedSize(280,320)


            oldPassword_label = QLabel("كلمة المرور القديمة",self.changePasswordWindow)
            oldPassword_label.setStyleSheet("font-size:25px")
            oldPassword_label.move(50,10)
            self.oldPassword = QLineEdit(self.changePasswordWindow)
            self.oldPassword.setStyleSheet("font-size:20px;width:240px")
            self.oldPassword.move(10,50)
            self.oldPassword.setEchoMode(QLineEdit.EchoMode.Password)

            newPassword_label1 = QLabel("كلمة المرور الجديدة",self.changePasswordWindow)
            newPassword_label1.setStyleSheet("font-size:25px")
            newPassword_label1.move(50,90)
            self.newPassword1 = QLineEdit(self.changePasswordWindow)
            self.newPassword1.setStyleSheet("font-size:20px;width:240px")
            self.newPassword1.move(10,130)
            self.newPassword1.setEchoMode(QLineEdit.EchoMode.Password)


            newPassword_label2 = QLabel("تأكيد كلمة المرور الجديدة",self.changePasswordWindow)
            newPassword_label2.setStyleSheet("font-size:25px")
            newPassword_label2.move(20,170)
            self.newPassword2 = QLineEdit(self.changePasswordWindow)
            self.newPassword2.setStyleSheet("font-size:20px;width:240px")
            self.newPassword2.move(10,210)
            self.newPassword2.setEchoMode(QLineEdit.EchoMode.Password)

            submitButton = QPushButton("تغيير",self.changePasswordWindow,clicked=self.changePasswordComplete)
            submitButton.setStyleSheet("font-size:20px;width:150px;background-color:green")
            submitButton.move(60,270)


            self.changePasswordWindow.show()

        def changePasswordComplete(self):
            cr.execute("SELECT password from loginStuff")
            oldPassword = cr.fetchone()[0]

            if oldPassword == self.oldPassword.text():
                if self.newPassword1.text()==self.newPassword2.text() and len(self.newPassword1.text()) > 0:
                    cr.execute(f"UPDATE loginStuff SET password = '{self.newPassword1.text()}'")
                    d = QMessageBox(parent =self.changePasswordWindow , text="تم تغيير كلمة المرور بنجاح")
                    d.setWindowTitle("نجاح")
                    d.setIcon(QMessageBox.Icon.Information)
                    d.setStyleSheet("background-color:white")
                    ret = d.exec()
                    con.commit()
                    self.changePasswordWindow.destroy()
                else:
                    msg = QMessageBox(parent=self,text="كلمات المرور غير متطابقات")
                    msg.setWindowTitle("ERROR")
                    msg.setStyleSheet("background-color:white")
                    msg.setIcon(QMessageBox.Icon.Critical)
                    msg.exec()
            else:
                d = QMessageBox(parent=self.changePasswordWindow,text="كلمة المرور خاطئة")
                d.setWindowTitle("ERROR")
                d.setIcon(QMessageBox.Icon.Critical)
                d.setStyleSheet("background-color:white")
                ret = d.exec()

        def developerFun(self):
            msg = QMessageBox(parent=self,text="م/عبدالله الشامي\nرقم الجوال:0558967920")
            msg.setStyleSheet("background-color:white")
            msg.setIcon(QMessageBox.Icon.Information)
            msg.setWindowTitle("مبرمج")
            msg.exec()
    class Window(QWidget):
        def __init__(self):
            super().__init__()
            self.confirmWindow = 0
            uic.loadUi("mainWindow.ui",self)
            self.setWindowTitle(title)
            self.setWindowIcon(QIcon(icon))
            self.resize(1120,680)
            self.appcentsDetalisWindow = 0
            self.employess.setColumnCount(8)
            self.employess.setHorizontalHeaderLabels(["الاسم","نوع الوظيفة","السجل المدني","الرقم الوظيفي","التخصص","المستوى والدرجة","الايميل","الراتب"])
            

            self.searchEntry.textChanged.connect(self.search)
            self.employess.setColumnWidth(0,80)
            self.employess.setColumnWidth(1,85)
            self.employess.setColumnWidth(2,110)
            self.employess.setColumnWidth(3,120)
            self.employess.setColumnWidth(4,120)
            self.employess.setColumnWidth(5,138)
            self.employess.setColumnWidth(6,80)
            self.employess.setColumnWidth(7,50)


            # self.employess.setStyleSheet("background-color:rgb(208, 255, 251)")
            self.exportDatabase.clicked.connect(self.exportDatabaseFun)
            self.insertDatabase.clicked.connect(self.insertDatabaseFun)

            self.scroll = QScrollArea()             # Scroll Area which contains the widgets, set as the centralWidget
            sop1.resizeEvent = self.resizeEvent
            self.pryW = self.width()
            self.pryH = self.height()


            #Scroll Area Properties
            self.scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOn)
            self.scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOn)
            # self.scroll.setWidgetResizable(True)
            self.scroll.setWidget(self)
            sop1.setCentralWidget(self.scroll)
            self.showComponents()
            # self.show()
            sop1.show()
        def resizeEvent(self, event):
            # Adjust the scroll area size to fit the window
            super().resizeEvent(event)
            



            # self.scroll.verticalScrollBar().setValue(self.scroll.verticalScrollBar().maximum())

            new_width = sop1.width()
            current_height = sop1.height() - 178
            if sop1.width() >=1120:
                self.resize(sop1.width(),sop1.height())
                self.currentW = self.width()
                self.currentH = self.height()
                # logo_Width = self.mainWindow.logo.width()
                # max_logo_position = new_width - logo_Width -5
                # self.logo_position = max_logo_position if new_width > frame_width else 900
                # self.mainWindow.logo.move(self.logo_position, 0)
                LabelNameWidth = self.HeaderPrograme.width()
                max_LabelNameWidth = new_width - LabelNameWidth - self.label.width() -350
                self.logo_position = max_LabelNameWidth if new_width > LabelNameWidth and max_LabelNameWidth > 250 else 250
                self.HeaderPrograme.move(self.logo_position, 10)

                LabelLogoWidth = self.label.width()
                max_LabelLogoWidth = new_width - LabelLogoWidth
                self.logoPhoto_position = max_LabelLogoWidth if new_width > LabelLogoWidth else 770
                self.label.move(self.logoPhoto_position, 0)

                exportButtonWidth = self.exportDatabase.width()
                max_exportButtonWidth = new_width - exportButtonWidth - LabelLogoWidth - 200
                self.exportButtonWidth_position = max_exportButtonWidth if new_width > exportButtonWidth and max_exportButtonWidth > 600 else 600
                self.exportDatabase.move(self.exportButtonWidth_position, 60)
                self.insertDatabase.move(self.exportButtonWidth_position,120)

                searchEntryWidth = self.frame_2.width()
                max_searchEntryWidth = new_width - exportButtonWidth - searchEntryWidth - 608
                self.searchEntry_position = max_searchEntryWidth if new_width > searchEntryWidth and max_searchEntryWidth > 10 else 10
                self.frame_2.move(self.searchEntry_position,70)

                # showFrameFakeShow
                a = new_width - self.showFrameFakeShow.width()
                b = self.height()- self.showFrameFakeShow.height()

                aa = new_width - self.frame.width() - 10
                self.frame_position = aa if new_width > self.frame.width() else 820
                self.frame.move(self.frame_position,self.frame.y())
                self.showFrameFakeShow.resize(self.showFrameFakeShow.width()+a,self.showFrameFakeShow.height()+b)
                event.accept()


                difrenceH = self.currentH- self.pryH
                difrenceW = self.currentW- self.pryW
                columnsWidths = ceil(difrenceW / 8)
                for i in range(self.employess.columnCount()):
                    self.employess.setColumnWidth(i,self.employess.columnWidth(i) + columnsWidths)

                self.employess.resize(self.employess.width() + difrenceW,self.employess.height() + difrenceH)
                self.pryW = self.width()
                self.pryH = self.height()
            else:
                if self.employess.columnCount() == 8:
                    regularWidth = [80,85,110,120,120,138,80,55]
                    for i in range(self.employess.columnCount()):
                        self.employess.setColumnWidth(i,regularWidth[i])
        def addemployee(self):
            for i in self.frame.children():
                i.deleteLater()
            self.frame.setGeometry(self.frame.x(),40,291,670)

            nameLabel = QLabel("الأسم",self.frame)
            nameLabel.setStyleSheet("background-color:#389D63;width:250px;color:black")
            nameLabel.setFont(QFont("Arial",18))
            nameLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            nameLabel.setGeometry(10,0,270,30)

            self.nameEntry = QLineEdit(self.frame)
            self.nameEntry.setFont(QFont("Arial",20))
            self.nameEntry.setGeometry(10,35,270,30)
            self.nameEntry.setStyleSheet("background-color:white")

            jopLabel = QLabel("نوع الوظيفة",self.frame)
            jopLabel.setStyleSheet("background-color:#389D63;width:250px;color:black")
            jopLabel.setFont(QFont("Arial",18))
            jopLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            jopLabel.setGeometry(10,70,270,30)

            self.jopEntry = QLineEdit(self.frame)
            self.jopEntry.setFont(QFont("Arial",20))
            self.jopEntry.setGeometry(10,105,270,30)
            self.jopEntry.setStyleSheet("background-color:white")

            reportLabel = QLabel("السجل المدني",self.frame)
            reportLabel.setStyleSheet("background-color:#389D63;width:250px;color:black")
            reportLabel.setFont(QFont("Arial",18))
            reportLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            reportLabel.setGeometry(10,140,270,30)

            self.reportEntry = QLineEdit(self.frame)
            self.reportEntry.setFont(QFont("Arial",20))
            self.reportEntry.setGeometry(10,175,270,30)
            self.reportEntry.setStyleSheet("background-color:white")

            numberLabel = QLabel("الرقم الوظيفي",self.frame)
            numberLabel.setStyleSheet("background-color:#389D63;width:250px;color:black")
            numberLabel.setFont(QFont("Arial",18))
            numberLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            numberLabel.setGeometry(10,210,270,30)

            self.numberEntry = QLineEdit(self.frame)
            self.numberEntry.setFont(QFont("Arial",20))
            self.numberEntry.setGeometry(10,245,270,30)
            self.numberEntry.setStyleSheet("background-color:white")

            spicialLabel = QLabel("التخصص",self.frame)
            spicialLabel.setStyleSheet("background-color:#389D63;width:250px;color:black")
            spicialLabel.setFont(QFont("Arial",18))
            spicialLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            spicialLabel.setGeometry(10,280,270,30)

            self.spicialEntry = QLineEdit(self.frame)
            self.spicialEntry.setFont(QFont("Arial",20))
            self.spicialEntry.setGeometry(10,315,270,30)
            self.spicialEntry.setStyleSheet("background-color:white")
            self.spicialEntry.setGeometry(10,315,270,30)


            leavellLabel = QLabel("المستوى والدرجة",self.frame)
            leavellLabel.setStyleSheet("background-color:#389D63;width:250px;color:black")
            leavellLabel.setFont(QFont("Arial",18))
            leavellLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            leavellLabel.setGeometry(10,350,270,30)

            self.leavelEntry = QLineEdit(self.frame)
            self.leavelEntry.setFont(QFont("Arial",20))
            self.leavelEntry.setGeometry(10,385,270,30)
            self.leavelEntry.setStyleSheet("background-color:white")

            EmailLabel = QLabel("الايميل",self.frame)
            EmailLabel.setStyleSheet("background-color:#389D63;width:250px;color:black")
            EmailLabel.setFont(QFont("Arial",16))
            EmailLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            EmailLabel.setGeometry(10,420,270,30)

            self.EmailEntry = QLineEdit(self.frame)
            self.EmailEntry.setFont(QFont("Arial",16))
            self.EmailEntry.setGeometry(10,455,270,30)
            self.EmailEntry.setStyleSheet("background-color:white")



            salaryLabel = QLabel("الراتب",self.frame)
            salaryLabel.setStyleSheet("background-color:#389D63;width:250px;color:black")
            salaryLabel.setFont(QFont("Arial",18))
            salaryLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            salaryLabel.setGeometry(10,490,270,30)

            self.salaryEntry = QLineEdit(self.frame)
            self.salaryEntry.setFont(QFont("Arial",20))
            self.salaryEntry.setGeometry(10,525,270,30)
            self.salaryEntry.setStyleSheet("background-color:white")

            self.salaryEntry.setText('0')

            completeAddButton = QPushButton("اضافة",self.frame)
            completeAddButton.setFont(QFont("Arial",20))
            completeAddButton.setGeometry(10,560,270,40)
            completeAddButton.setStyleSheet("background-color:#389D63;color:black")

            completeAddButton.clicked.connect(self.completeAddemployee)

            backAddButton = QPushButton("العودة",self.frame,clicked=self.showComponents)
            backAddButton.setFont(QFont("Arial",20))
            backAddButton.setGeometry(15,605,260,40)
            backAddButton.setStyleSheet("background-color:#389D63;color:black")

            EmailLabel.show()
            self.EmailEntry.show()
            salaryLabel.show()
            self.salaryEntry.show()
            nameLabel.show()
            self.nameEntry.show()
            jopLabel.show()
            self.jopEntry.show()
            reportLabel.show()
            self.reportEntry.show()
            numberLabel.show()
            self.numberEntry.show()
            spicialLabel.show()
            self.spicialEntry.show()
            leavellLabel.show()
            self.leavelEntry.show()
            completeAddButton.show()
            backAddButton.show()
        def completeAddemployee(self):

            if len(self.nameEntry.text()) > 0 and len(self.jopEntry.text()) > 0 and len(self.reportEntry.text()) > 0 and len(self.numberEntry.text()) > 0 and len(self.spicialEntry.text()) > 0 and len(self.leavelEntry.text()) > 0 and len(self.salaryEntry.text()) > 0 and len(self.EmailEntry.text()) > 0:
                number = self.reportEntry.text()
                cr.execute(f"SELECT * FROM info WHERE identy='{str(number)}'")
                a = cr.fetchall()
                if len(a) == 0:
                    try:
                        int(self.salaryEntry.text())
                        cr.execute(f"INSERT INTO info (name,jop,identy,numberJop,spicialLabel,leavelLabel,email,salary) values ('{self.nameEntry.text()}','{self.jopEntry.text()}','{self.reportEntry.text()}','{self.numberEntry.text()}','{self.spicialEntry.text()}','{self.leavelEntry.text()}','{self.EmailEntry.text()}','{self.salaryEntry.text()}')")
                        con.commit()
                        d = QMessageBox(parent=self,text="تم اضافة المتطوع بنجاح")
                        d.setWindowTitle("نجاح")
                        d.setIcon(QMessageBox.Icon.Information)
                        d.exec()
                        self.updateView()
                    except:
                        d = QMessageBox(parent=self,text="يرجى كتابة الراتب رقما فقط")
                        d.setWindowTitle("ERROR")
                        d.setIcon(QMessageBox.Icon.Critical)
                        d.exec()
                else:
                    d = QMessageBox(parent=self,text="هناك متطوع بنفس رقم السجل المدني")
                    d.setWindowTitle("ERROR")
                    d.setIcon(QMessageBox.Icon.Critical)
                    d.exec()
            else:
                d = QMessageBox(parent=self,text="يجب ملئ جميع الحقول")
                d.setWindowTitle("ERROR")
                d.setIcon(QMessageBox.Icon.Critical)
                d.exec()
        def showComponents(self):
            self.searchEntry.setDisabled(False)
            
            try:
                self.frame.setGeometry(self.frame.x(),190,291,481)
            except:
                self.frame = QFrame(self)
                self.frame.setGeometry(820,190,291,481)
                self.frame.setStyleSheet("background-color:rgb(0, 152, 0)")
            for i in self.frame.children():
                i.deleteLater()

            label = QLabel("الرئيسيه",self.frame)
            label.setStyleSheet("background-color:#44C17A;width:250px;color:black")
            label.setFont(QFont("Arial",20))
            label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            label.setGeometry(0,0,311,25)

            add_button = QPushButton("اضافة متطوع",self.frame)
            add_button.setStyleSheet("background-color:#44C17A;width:250px;color:black")
            add_button.setFont(QFont("Arial",20))
            add_button.move(17,40)

            delete_button = QPushButton("حذف متطوع",self.frame)
            delete_button.setStyleSheet("background-color:#44C17A;width:250px;color:black")
            delete_button.setFont(QFont("Arial",20)) 
            delete_button.move(17,78)

            appsent_button = QPushButton("تسجيل غياب متطوع",self.frame,clicked=self.setAppcent)
            appsent_button.setStyleSheet("background-color:#44C17A;width:250px;color:black")
            appsent_button.setFont(QFont("Arial",20))
            appsent_button.move(17,116)

            edit_button = QPushButton("تعديل معلومات متطوع",self.frame,clicked=self.editInfo)
            edit_button.setStyleSheet("background-color:#44C17A;width:250px;color:black")
            edit_button.setFont(QFont("Arial",20))
            edit_button.move(17,154)

            exportOne_button = QPushButton("تصدير غياب متطوع معين",self.frame,clicked=self.exportOne)
            exportOne_button.setStyleSheet("background-color:#44C17A;width:250px;color:black")
            exportOne_button.setFont(QFont("Arial",20))
            exportOne_button.move(17,192)

            exportAll_button = QPushButton("تصدير الغياب لجميع المتطوعين",self.frame,clicked=self.exportAll)
            exportAll_button.setStyleSheet("background-color:#44C17A;width:250px;color:black")
            exportAll_button.setFont(QFont("Arial",18))
            exportAll_button.move(17,230)

            IssueAdd_button = QPushButton("اضافة مسائلة لمتطوع",self.frame,clicked=self.addIssueEmployee)
            IssueAdd_button.setStyleSheet("background-color:#44C17A;width:250px;color:black")
            IssueAdd_button.setFont(QFont("Arial",20))
            IssueAdd_button.move(17,265)

            exportOneIssues_button = QPushButton("تصدير المسائلات لمتطوع",self.frame,clicked=self.exportSingleIssues)
            exportOneIssues_button.setStyleSheet("background-color:#44C17A;width:250px;color:black")
            exportOneIssues_button.setFont(QFont("Arial",20))
            exportOneIssues_button.move(17,303)

            exportAllIssues_button = QPushButton("تصدير المسائلات لجميع المتطوعين",self.frame,clicked=self.exportAllIssues)
            exportAllIssues_button.setStyleSheet("background-color:#44C17A;width:250px;color:black")
            exportAllIssues_button.setFont(QFont("Arial",15))
            exportAllIssues_button.move(17,341)

            exportSalary_button = QPushButton("تصدير كشف الرواتب",self.frame,clicked=self.exportSalaryReport)
            exportSalary_button.setStyleSheet("background-color:#44C17A;width:250px;color:black")
            exportSalary_button.setFont(QFont("Arial",17))
            exportSalary_button.move(17,371)

            send_button = QPushButton("ارسال الرسائل",self.frame,clicked=self.sendMessages)
            send_button.setStyleSheet("background-color:#44C17A;width:250px;color:black")
            send_button.setFont(QFont("Arial",17))
            send_button.move(17,405)


            add_button.clicked.connect(self.addemployee)
            delete_button.clicked.connect(self.deleteEmployee)
            self.updateView()

            send_button.show()
            label.show()
            add_button.show()
            delete_button.show()
            appsent_button.show()
            edit_button.show()
            exportOne_button.show()
            exportAll_button.show()
            IssueAdd_button.show()
            exportOneIssues_button.show()
            exportAllIssues_button.show()
            exportSalary_button.show()
            self.frame.show()
        def updateView(self):
            self.employess.setRowCount(0)
            if self.employess.columnCount() >=9:
                self.employess.removeColumn(0)
            cr.execute("SELECT * FROM info")
            tempThing = [] 
            for i in cr.fetchall():
                tempThing.append(i)
            for row,i in enumerate(tempThing):
                self.employess.insertRow(self.employess.rowCount())
                for col,val in enumerate(i):
                    self.employess.setItem(row,col,QTableWidgetItem(str(val)))
            for row in range(self.employess.rowCount()):
                for col in range(self.employess.columnCount()):
                    self.employess.item(row,col).setFlags(Qt.ItemFlag.ItemIsEditable)
        def deleteEmployee(self):
            for i in self.frame.children():
                i.deleteLater()
            self.frame.setGeometry(self.frame.x(),190,291,111)
            self.frame.setStyleSheet("background-color:rgb(0, 152, 0)")
            

            completeDeleteButton = QPushButton("حذف",self.frame,clicked=self.completeDeleteEmployee)
            completeDeleteButton.setFont(QFont("Arial",20))
            completeDeleteButton.setGeometry(10,10,270,40)
            completeDeleteButton.setStyleSheet("background-color:red;color:white")

            CancelDeleteButton = QPushButton("الغاء",self.frame)
            CancelDeleteButton.setFont(QFont("Arial",20))
            CancelDeleteButton.setGeometry(10,60,270,40)
            CancelDeleteButton.setStyleSheet("background-color:green;color:white")
            CancelDeleteButton.clicked.connect(self.showComponents)

            self.TableToDelete()
            completeDeleteButton.show()
            CancelDeleteButton.show()
            self.frame.show()
        def TableToDelete(self):
            self.searchEntry.setDisabled(True)
            for row in range(self.employess.rowCount()):
                item = QTableWidgetItem(self.employess.item(row,0).text())
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable|Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(Qt.CheckState.Unchecked)
                self.employess.setItem(row,0,item)

            for row in range(self.employess.rowCount()):
                self.employess.item(row,1).setFlags(Qt.ItemFlag.ItemIsEditable)
                self.employess.item(row,2).setFlags(Qt.ItemFlag.ItemIsEditable)
                self.employess.item(row,3).setFlags(Qt.ItemFlag.ItemIsEditable)
                self.employess.item(row,4).setFlags(Qt.ItemFlag.ItemIsEditable)
                self.employess.item(row,5).setFlags(Qt.ItemFlag.ItemIsEditable)
        def completeDeleteEmployee(self):
            numbers = []
            Names = []
            for row in range(self.employess.rowCount()):
                if self.employess.item(row,0).checkState()==Qt.CheckState.Checked:
                    numbers.append(self.employess.item(row,2).text())
                    Names.append(self.employess.item(row,0).text())
                    Names.append("\n")
            
            d = QMessageBox(parent=self,text=f"تأكيد حذف {len(numbers)} متطوع ")
            d.setIcon(QMessageBox.Icon.Information)
            d.setWindowTitle("نجاح")
            d.setStandardButtons(QMessageBox.StandardButton.Cancel|QMessageBox.StandardButton.Ok)
            d.setDetailedText("".join(Names))
            important = d.exec()
            if important == QMessageBox.StandardButton.Ok:
                for i in numbers:
                    cr.execute(f"SELECT id FROM vactions WHERE employeeId='{i}'")
                    for exza in cr.fetchall():
                        for j in exza:
                            cr.execute(f"DELETE FROM images WHERE vactionId='{j}'")

                    cr.execute(f"SELECT id FROM issues WHERE employeeIdenty='{i}'")
                    for exza in cr.fetchall():
                        for j in exza:
                            cr.execute(f"DELETE FROM imagesIssues WHERE issuesId='{j}'")

                    cr.execute(f"DELETE FROM info WHERE identy='{i}'")
                    cr.execute(f"DELETE FROM issues WHERE employeeIdenty='{i}'")
                    cr.execute(f"DELETE FROM vactions WHERE employeeId='{i}'")

                con.commit()
                d = QMessageBox(parent=self,text="تم الحذف بنجاح")
                
                d.setIcon(QMessageBox.Icon.Information)
                d.setWindowTitle("نجاح")
                d.exec()
            self.showComponents()
        def editInfo(self):
            for i in self.frame.children():
                i.deleteLater()
            # self.frame.setGeometry(820,300,291,610) # 820,190,291,521
            self.frame.setGeometry(self.frame.x(),0,291,670)
            self.frame.setStyleSheet("background-color:rgb(0, 152, 0)")

            self.searchEntry.setDisabled(True)

            nameLabel = QLabel("الأسم",self.frame)
            nameLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            nameLabel.setFont(QFont("Arial",18))
            nameLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            nameLabel.setGeometry(10,0,270,30)

            self.nameEntry = QLineEdit(self.frame)
            self.nameEntry.setFont(QFont("Arial",20))
            self.nameEntry.setGeometry(10,35,270,30)
            self.nameEntry.setStyleSheet("background-color:white")

            jopLabel = QLabel("نوع الوظيفة",self.frame)
            jopLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            jopLabel.setFont(QFont("Arial",18))
            jopLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            jopLabel.setGeometry(10,70,270,30)
            
            self.jopEntry = QLineEdit(self.frame)
            self.jopEntry.setFont(QFont("Arial",20))
            self.jopEntry.setGeometry(10,105,270,30)
            self.jopEntry.setStyleSheet("background-color:white")
            
            reportLabel = QLabel("السجل المدني",self.frame)
            reportLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            reportLabel.setFont(QFont("Arial",18))
            reportLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            reportLabel.setGeometry(10,140,270,30)

            self.reportEntry = QLineEdit(self.frame)
            self.reportEntry.setFont(QFont("Arial",20))
            self.reportEntry.setGeometry(10,175,270,30)
            self.reportEntry.setStyleSheet("background-color:white")

            numberLabel = QLabel("الرقم الوظيفي",self.frame)
            numberLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            numberLabel.setFont(QFont("Arial",18))
            numberLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            numberLabel.setGeometry(10,210,270,30)

            self.numberEntry = QLineEdit(self.frame)
            self.numberEntry.setFont(QFont("Arial",20))
            self.numberEntry.setGeometry(10,245,270,30)
            self.numberEntry.setStyleSheet("background-color:white")

            spicialLabel = QLabel("التخصص",self.frame)
            spicialLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            spicialLabel.setFont(QFont("Arial",18))
            spicialLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            spicialLabel.setGeometry(10,280,270,30)

            self.spicialEntry = QLineEdit(self.frame)
            self.spicialEntry.setFont(QFont("Arial",20))
            self.spicialEntry.setGeometry(10,315,270,30)
            self.spicialEntry.setStyleSheet("background-color:white")
            self.spicialEntry.setGeometry(10,315,270,30)


            leavellLabel = QLabel("المستوى والدرجة",self.frame)
            leavellLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            leavellLabel.setFont(QFont("Arial",18))
            leavellLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            leavellLabel.setGeometry(10,350,270,30)

            self.leavelEntry = QLineEdit(self.frame)
            self.leavelEntry.setFont(QFont("Arial",20))
            self.leavelEntry.setGeometry(10,385,270,30)
            self.leavelEntry.setStyleSheet("background-color:white")

            emailLabel = QLabel("الايميل",self.frame)
            emailLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            emailLabel.setFont(QFont("Arial",16))
            emailLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            emailLabel.setGeometry(10,420,270,20)

            self.EmailEntry = QLineEdit(self.frame)
            self.EmailEntry.setFont(QFont("Arial",17))
            self.EmailEntry.setGeometry(10,445,270,25)
            self.EmailEntry.setStyleSheet("background-color:white")

            SalaryLabel = QLabel("الراتب",self.frame)
            SalaryLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            SalaryLabel.setFont(QFont("Arial",18))
            SalaryLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            SalaryLabel.setGeometry(10,475,270,25)

            self.salaryEntry = QLineEdit(self.frame)
            self.salaryEntry.setFont(QFont("Arial",20))
            self.salaryEntry.setGeometry(10,510,270,25)
            self.salaryEntry.setStyleSheet("background-color:white")


            self.IssuesButton = QPushButton("تفاصيل المسائلات",self.frame,clicked=self.issuesDetalis)
            self.IssuesButton.setFont(QFont("Arial",20))
            self.IssuesButton.setGeometry(10,545,270,30)
            self.IssuesButton.setStyleSheet("background-color:#389D63;color:white")


            self.detalisButton = QPushButton("تفاصيل الغياب",self.frame,clicked=self.appcentsDetalis)
            self.detalisButton.setFont(QFont("Arial",20))
            self.detalisButton.setGeometry(10,580,270,25)
            self.detalisButton.setStyleSheet("background-color:#389D63;color:white")
            self.detalisButton.setDisabled(True)



            completeAddButton = QPushButton("تعديل",self.frame)
            completeAddButton.setFont(QFont("Arial",20))
            completeAddButton.setGeometry(10,610,270,25)
            completeAddButton.setStyleSheet("background-color:#389D63;color:white")

            completeAddButton.clicked.connect(self.completeEdit)

            backAddButton = QPushButton("العودة",self.frame,clicked=self.showComponents)
            backAddButton.setFont(QFont("Arial",16))
            backAddButton.setGeometry(15,640,260,25)
            backAddButton.setStyleSheet("background-color:#389D63;color:white")
            

            self.showRadioButtonsToEdit("Edit")
            emailLabel.show()
            self.EmailEntry.show()
            SalaryLabel.show()
            self.salaryEntry.show()
            nameLabel.show()
            self.nameEntry.show()
            jopLabel.show()
            self.jopEntry.show()
            reportLabel.show()
            self.reportEntry.show()
            numberLabel.show()
            self.numberEntry.show()
            spicialLabel.show()
            self.spicialEntry.show()
            leavellLabel.show()
            self.leavelEntry.show()
            self.IssuesButton.show()
            self.detalisButton.show()
            completeAddButton.show()
            backAddButton.show()
            self.frame.show()
        def appcentsDetalis(self):
            self.delatedItems = []
            
            self.appcentsDetalisWindow = tempwindows()
            self.appcentsDetalisWindow.setFixedSize(800,600)
            self.appcentsDetalisTable = QTableWidget(self.appcentsDetalisWindow)
            self.appcentsDetalisTable.setGeometry(10,40,700,400)
            self.appcentsDetalisTable.setColumnCount(13) 
            self.appcentsDetalisTable.setColumnHidden(11,True)
            self.appcentsDetalisTable.setColumnWidth(0,30)
            self.appcentsDetalisTable.setHorizontalHeaderLabels(["","اسم المتطوع","نوع الاجازة","اليوم","حالة الاجازة","نظام فارس","ملاحظات","نوع التقرير","جهة التقرير","التاريخ بالهجري","التاريخ بالميلادي","","صورة العذر"])
            self.appcentsDetalisTable.setColumnHidden(5,True)
            self.appcentsDetalisTable.setColumnHidden(7,True)
            self.appcentsDetalisTable.setColumnHidden(8,True)


            buttonSaveEveryThing = QPushButton(self.appcentsDetalisWindow,clicked=self.comleteEveryThingINAppcentTable)
            buttonSaveEveryThing.setStyleSheet("background-color:green;color:white")
            buttonSaveEveryThing.setText("حفظ التغيرات")
            buttonSaveEveryThing.setGeometry(320,450,150,40)

            cr.execute(f"SELECT * FROM vactions WHERE employeeId='{self.reportEntry.text()}'")
            infos = []
            for i in cr.fetchall():
                listi = list(i)
                cr.execute(f"SELECT name FROM info WHERE identy='{listi[0]}'")
                listi[0] = cr.fetchone()[0]
                listi[-1] = str(listi[-1])
                listi.insert(0,"")
                infos.append(listi)
            for row,i in enumerate(infos):
                self.appcentsDetalisTable.insertRow(self.appcentsDetalisTable.rowCount())
                for col,val in enumerate(i):
                    if col%12==0:
                        icon = QPixmap("assests/trashicon.png")
                        button = QPushButton()
                        button.setStyleSheet(f"Qproperty-icon:url(assests/trashicon.png);qproperty-iconSize:30px 30px;background-color:rgb(253, 253, 253)")
                        button.clicked.connect(lambda x,row=row:self.deleteFromAppcentTable(row))
                        self.appcentsDetalisTable.setIndexWidget(self.appcentsDetalisTable.model().index(row,0),button)
                    else:
                        self.appcentsDetalisTable.setItem(row,col,QTableWidgetItem(val))
            for i in range(self.appcentsDetalisTable.rowCount()):
                self.appcentsDetalisTable.item(i,10).setFlags(Qt.ItemFlag.ItemIsEditable)
                self.appcentsDetalisTable.item(i,9).setFlags(Qt.ItemFlag.ItemIsEditable)
                self.appcentsDetalisTable.item(i,1).setFlags(Qt.ItemFlag.ItemIsEditable)
            
            for i in range(self.appcentsDetalisTable.rowCount()):
                cr.execute(f"SELECT img,state FROM images where vactionId='{self.appcentsDetalisTable.item(i,11).text()}'")
                imagesData = cr.fetchone()
                if imagesData !=None:
                    item = self.returnQLabelWithImage(imagesData[0],imagesData[1])
                    self.appcentsDetalisTable.setCellWidget(i,12,item)
            self.appcentsDetalisTable.cellDoubleClicked.connect(self.insertImageToACell)
            # cant append to images to one execuse
            self.appcentsDetalisWindow.show()
        def insertImageToACell(self,row,column):
            if column==12:
                cr.execute(f"DELETE FROM images WHERE vactionId = '{str(self.appcentsDetalisTable.item(row,11).text())}'")
                fileImageUser = QFileDialog.getOpenFileName(parent=self.appcentsDetalisWindow,caption="Select a File",filter="Image File (*.png *.jpg *.jpeg)")
                typeFile = os.path.splitext(fileImageUser[0])[1]
                if len(fileImageUser[0]) > 0:
                    with open(fileImageUser[0],"rb") as binary_image:
                        binaryCode = binary_image.read()
                    cr.execute("""INSERT INTO images (img,vactionId,state) values(?,?,?)""",(binaryCode,str(self.appcentsDetalisTable.item(row,11).text()),typeFile))
                    item = self.returnQLabelWithImage(binaryCode,typeFile)
                    self.appcentsDetalisTable.setCellWidget(row,column,item)
        def deleteFromAppcentTable(self,row):
            self.delatedItems.append(self.appcentsDetalisTable.item(row,11).text())
            # self.appcentsDetalisTable.removeRow(row) -> cause an error
            self.appcentsDetalisTable.hideRow(row) # -> solve the error # -> it cause another problem
        def comleteEveryThingINAppcentTable(self):

            NativeResult = []

            for row in range(self.appcentsDetalisTable.rowCount()):
                NativeResult.append([])
                for col in range(self.appcentsDetalisTable.columnCount()-1):
                    if col!=0:
                        NativeResult[row].append(self.appcentsDetalisTable.item(row,col).text())
                NativeResult[row][0] = self.reportEntry.text()
            cr.execute(f"SELECT * FROM vactions WHERE employeeId='{self.reportEntry.text()}'")
            databaseResult = []
            for i in cr.fetchall():
                listiThing = list(i)
                listiThing[-1] = str(listiThing[-1])
                databaseResult.append(list(listiThing))

            for i in range(len(NativeResult)):
                if NativeResult[i] != databaseResult[i]:
                    cr.execute(f"UPDATE vactions set vactionType='{NativeResult[i][1]}',day='{NativeResult[i][2]}',execute='{NativeResult[i][3]}',Fares='{NativeResult[i][4]}',notes='{NativeResult[i][5]}',whereFrom='{NativeResult[i][6]}',whereExact='{NativeResult[i][7]}',HijryDate='{NativeResult[i][8]}',GeoDate='{NativeResult[i][9]}' WHERE id='{NativeResult[i][-1]}'")
            for i in self.delatedItems:
                cr.execute(f"DELETE FROM vactions WHERE id='{i}'")
                cr.execute(f"DELETE FROM images WHERE vactionId='{i}'")

            self.appcentsDetalisWindow.destroy()
            self.appcentsDetalisWindow.hide()

            con.commit()
        def completeEdit(self):
            if len(self.nameEntry.text()) > 0 and len(self.jopEntry.text()) > 0 and len(self.reportEntry.text()) > 0 and len(self.numberEntry.text()) > 0 and len(self.spicialEntry.text()) > 0 and len(self.leavelEntry.text()) > 0 and len(self.salaryEntry.text()) > 0 and len(self.EmailEntry.text()) > 0:
                cr.execute(f"Update info SET name ='{self.nameEntry.text()}' , jop ='{self.jopEntry.text()}' , numberJop='{self.numberEntry.text()}',spicialLabel ='{self.spicialEntry.text()}' ,leavelLabel ='{self.leavelEntry.text()}',email='{self.EmailEntry.text()}',salary='{self.salaryEntry.text()}' WHERE identy='{self.reportEntry.text()}'")
                cr.execute(f"UPDATE issues SET employeeName='{self.nameEntry.text()}' WHERE employeeIdenty='{self.reportEntry.text()}'")
                con.commit()
                d = QMessageBox(parent=self,text="تم تعديل معلومات المتطوع بنجاح")
                d.setWindowTitle("نجاح")
                d.setIcon(QMessageBox.Icon.Information)
                d.exec()
                self.showComponents()
        def returnQLabelWithImage(self,binCode,state):
            imageLabel = QLabel(self.appcentsDetalisWindow)
            imageLabel.setScaledContents(True)
            pixmap = QPixmap()
            pixmap.loadFromData(binCode,state)
            imageLabel.setPixmap(pixmap)
            return imageLabel
        def showRadioButtonsToEdit(self,Where):
            # self.employess.setRowCount(0)

            # cr.execute("SELECT * FROM info")
            # tempThing = [] 
            self.employess.insertColumn(0)
            self.employess.setHorizontalHeaderLabels(["","الاسم","نوع الوظيفة","السجل المدني","الرقم الوظيفي","التخصص","المستوى والدرجة"])
            self.employess.setColumnWidth(0,20)
            self.searchEntry.setDisabled(True)
            
            # for e,i in enumerate(cr.fetchall()):
            #     tempThing.append(list(i))
            #     tempThing[e].insert(0,"")
            for row in range(self.employess.rowCount()):
                # self.employess.insertRow(self.employess.rowCount())
                # for col,val in enumerate(i):
                    # if col%7==0:
                        button = QRadioButton()
                        button.clicked.connect(lambda ch,o=row,Where=Where:self.transferInfo(o,Where))
                        self.employess.setIndexWidget(self.employess.model().index(row,0),button)
                    # else:
                    #     self.employess.setItem(row,col,QTableWidgetItem(val))

            for row in range(self.employess.rowCount()):
                    self.employess.item(row,1).setFlags(Qt.ItemFlag.ItemIsEditable)
                    self.employess.item(row,2).setFlags(Qt.ItemFlag.ItemIsEditable)
                    self.employess.item(row,3).setFlags(Qt.ItemFlag.ItemIsEditable)
                    self.employess.item(row,4).setFlags(Qt.ItemFlag.ItemIsEditable)
                    self.employess.item(row,5).setFlags(Qt.ItemFlag.ItemIsEditable)
        def transferInfo(self,row,Where):
            if Where == "Edit":
                self.nameEntry.setText(self.employess.item(row,1).text())
                self.jopEntry.setText(self.employess.item(row,2).text())
                self.reportEntry.setText(self.employess.item(row,3).text())
                self.numberEntry.setText(self.employess.item(row,4).text())
                self.spicialEntry.setText(self.employess.item(row,5).text())
                self.leavelEntry.setText(self.employess.item(row,6).text())
                self.EmailEntry.setText(self.employess.item(row,7).text())
                self.salaryEntry.setText(self.employess.item(row,8).text())
                self.reportEntry.setDisabled(True)
                self.detalisButton.setDisabled(False)
            if Where =="Appcent":
                self.identyEntry.setText(self.employess.item(row,3).text())
                self.identyEntry.setDisabled(True)
        def search(self,text):
            if len(self.searchEntry.text())==0:
                self.showComponents()
            else:
                self.employess.setRowCount(0)
                if self.choice.currentText()=="السجل المدني":
                    cr.execute("SELECT identy from info")
                    choices = cr.fetchall()
                    posiple = []
                    for o in choices:
                        for n,i in enumerate(o):
                            try:
                                if o[n][:len(self.searchEntry.text())]==self.searchEntry.text():
                                        if i not in posiple: # -> added new it was before without this line the (if line)
                                            posiple.append(i)
                            except:
                                pass
                    for p in posiple:
                        cr.execute(f"SELECT * from info where identy='{p}'")
                        self.employess.insertRow(self.employess.rowCount())
                        for i in cr.fetchall():
                            for col,j in enumerate(i):
                                self.employess.setItem(self.employess.rowCount()-1,col,QTableWidgetItem(str(j)))
                else:
                    cr.execute("SELECT name from info")
                    choices = cr.fetchall()
                    posiple = []
                    for o in choices:
                        for n,i in enumerate(o):
                            x = (str(o[0])).split()
                            y = self.searchEntry.text().split()
                            try:
                                # if self.search_entry.get() in x:
                                #         posiple.append(i)
                                l = len(y)
                                for n in (x):
                                    for Q in (y):
                                        if Q == n[0:len(Q)]:
                                            l-=1
                                if l==0:
                                    if i not in posiple:# -> added new it was before without this line the (if line)
                                        posiple.append(i)
                            except:
                                pass
                    for p in posiple:
                        cr.execute(f"SELECT * from info where name='{p}'")
                        self.employess.insertRow(self.employess.rowCount())
                        for i in cr.fetchall():
                            for col,j in enumerate(i):
                                self.employess.setItem(self.employess.rowCount()-1,col,QTableWidgetItem(str(j)))
                for row in range(self.employess.rowCount()):
                    for col in range(self.employess.columnCount()):
                        self.employess.item(row,col).setFlags(Qt.ItemFlag.ItemIsEditable)
        def setAppcent(self):
            # self.frame.setGeometry(820,210,291,694) # 820,120,291,590
            for i in self.frame.children():
                i.deleteLater()
            self.frame.setGeometry(self.frame.x(),40,291,620)
            self.addtional = False
            self.otherVaction = False
            
            identyLabel = QLabel("السجل المدني",self.frame)
            identyLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            identyLabel.setFont(QFont("Arial",20))
            identyLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            identyLabel.setGeometry(10,0,270,30)
            
            self.identyEntry = QLineEdit(self.frame)
            self.identyEntry.setFont(QFont("Arial",18))
            self.identyEntry.setGeometry(10,35,270,30)
            self.identyEntry.setStyleSheet("background-color:white")
            self.identyEntry.setDisabled(True)

            nameLabel = QLabel("الاجازة",self.frame)
            nameLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            nameLabel.setFont(QFont("Arial",20))
            nameLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            nameLabel.setGeometry(10,70,270,30)

            self.vications = QComboBox(self.frame)
            self.vications.setFont(QFont("Arial",15))
            self.vications.setGeometry(10,105,270,30)
            self.vications.setStyleSheet("background-color:white")

            self.vications.addItems(["اضطراري","مرضي","مرافق","دورة","لقاء","مراجعة","حافز","حالة وفاة","أخرى"])
            self.vications.activated.connect(self.addtionalThings)
            self.showRadioButtonsToEdit("Appcent")

            daylabel = QLabel("اليوم",self.frame)
            daylabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            daylabel.setFont(QFont("Arial",20))
            daylabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            daylabel.setGeometry(10,140,270,30)

            self.dayEntry = QLineEdit(self.frame)
            self.dayEntry.setFont(QFont("Arial",18))
            self.dayEntry.setGeometry(10,175,270,30)
            self.dayEntry.setStyleSheet("background-color:white")

            dateLabel = QLabel("التاريخ الميلادي",self.frame)
            dateLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            dateLabel.setFont(QFont("Arial",20))
            dateLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            dateLabel.setGeometry(10,210,270,30)


            self.dateEntry = QDateEdit(self.frame)
            self.dateEntry.setCalendarPopup(True)
            self.dateEntry.setDisplayFormat("yyyy/MM/dd")
            arabic_locale = QLocale(QLocale.Language.Arabic, QLocale.Country.SaudiArabia)
            self.dateEntry.setLocale(arabic_locale)
            self.dateEntry.setFont(QFont("Arial",18))
            self.dateEntry.setGeometry(10,245,270,30)
            self.dateEntry.setStyleSheet("background-color:white;color:black")
            self.todayButton = QPushButton("اليوم",clicked=lambda:self.dateEntry.calendarWidget().setSelectedDate(QDate().currentDate()))
            self.todayButton.setStyleSheet("background-color:green;")
            self.dateEntry.calendarWidget().layout().addWidget(self.todayButton)
            self.dateEntry.dateChanged.connect(self.changeHijryDate)


            dateHigryLabel = QLabel("التاريخ الهجري",self.frame)
            dateHigryLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            dateHigryLabel.setFont(QFont("Arial",20))
            dateHigryLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            dateHigryLabel.setGeometry(10,280,270,30)

            self.dateHigryEntry = QLineEdit(self.frame)
            self.dateHigryEntry.setFont(QFont("Arial",18))
            self.dateHigryEntry.setGeometry(10,315,270,30)
            self.dateHigryEntry.setStyleSheet("background-color:white")
            self.dateHigryEntry.setDisabled(True)

            self.withExecuse = QCheckBox('بعذر',self.frame)
            self.withExecuse.setGeometry(220,350,120,25)
            self.withExecuse.setFont(QFont("Arial",15))
            self.withExecuse.clicked.connect(lambda ch,x=0:self.Change(x))

            self.withOutExecuse = QCheckBox('بدون عذر',self.frame)
            self.withOutExecuse.setGeometry(30,350,120,25)
            self.withOutExecuse.setFont(QFont("Arial",15))
            self.withOutExecuse.clicked.connect(lambda ch,x=1:self.Change(x))

            NotesLabel = QLabel("الملاحظات",self.frame)
            NotesLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            NotesLabel.setFont(QFont("Arial",20))
            NotesLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            NotesLabel.setGeometry(10,380,270,25)

            self.NotesEntry = QLineEdit(self.frame)
            self.NotesEntry.setFont(QFont("Arial",18))
            self.NotesEntry.setGeometry(10,410,270,25)
            self.NotesEntry.setStyleSheet("background-color:white")
            
            completeAddAppcentButton = QPushButton("اضافة",self.frame)
            completeAddAppcentButton.setFont(QFont("Arial",20))
            completeAddAppcentButton.setGeometry(10,550,270,25)
            completeAddAppcentButton.setStyleSheet("background-color:#389D63;color:white")

            completeAddAppcentButton.clicked.connect(self.completeAddAppcent)

            backButton = QPushButton("العودة",self.frame,clicked=self.showComponents)
            backButton.setFont(QFont("Arial",20))
            backButton.setGeometry(15,585,260,30)
            backButton.setStyleSheet("background-color:#389D63;color:white")

            identyLabel.show()
            self.identyEntry.show()
            nameLabel.show()
            self.vications.show()
            self.vications.show()
            daylabel.show()
            self.dayEntry.show()
            dateLabel.show()
            self.dateEntry.show()
            dateHigryLabel.show()
            self.dateHigryEntry.show()
            self.withExecuse.show()
            self.withOutExecuse.show()
            NotesLabel.show()
            self.NotesEntry.show()
            completeAddAppcentButton.show()
            backButton.show()
            self.frame.show()
        def changeHijryDate(self):
            currentDate = (str(self.dateEntry.text()))
            currentDateHijry = (currentDate).split("/")
            currentDateHijryFinesh = str(Gregorian(int(currentDateHijry[0]),int(currentDateHijry[1]),int(currentDateHijry[2])).to_hijri()).replace("-","/")
            self.dateHigryEntry.setText(currentDateHijryFinesh)
        def completeAddAppcent(self):
            if len(self.identyEntry.text()) > 0 and len(self.vications.currentText()) > 0 and len(self.dayEntry.text()) >  0 and len(self.dateEntry.text()) >  0  and len(self.dateHigryEntry.text()) >  0 and (self.withExecuse.checkState()==Qt.CheckState.Checked or self.withOutExecuse.checkState()==Qt.CheckState.Checked):
                cr.execute(f"SELECT employeeId FROM vactions WHERE GeoDate = '{self.dateEntry.text()}' and employeeId='{self.identyEntry.text()}'")
                dates = cr.fetchall()
                if len(dates) > 0:
                    d = QMessageBox(parent=self,text="يوجد بالفعل عذر بنفس التاريخ هل تريد اضافة عذر اخر؟")
                    d.setWindowTitle("تأكيد")
                    d.setIcon(QMessageBox.Icon.Information)
                    d.setStandardButtons(QMessageBox.StandardButton.Cancel|QMessageBox.StandardButton.Ok)
                    state = d.exec()
                    if state == QMessageBox.StandardButton.Ok:
                        if self.withExecuse.checkState()==Qt.CheckState.Checked:
                            ExecuseThing = "بعذر"
                        elif self.withOutExecuse.checkState()==Qt.CheckState.Checked:
                            ExecuseThing = "بدون بعذر"
                        if self.addtional:
                            if (self.withFares.checkState()==Qt.CheckState.Checked or self.withOutFares.checkState()==Qt.CheckState.Checked) and (self.security.checkState()==Qt.CheckState.Checked or self.private.checkState()==Qt.CheckState.Checked) and (len(self.WhereEntry.text()) > 0):
                                if self.withFares.checkState()==Qt.CheckState.Checked:
                                    FaresThing = "تم"
                                elif self.withOutFares.checkState()==Qt.CheckState.Checked:
                                    FaresThing = "لم يتم"

                                if self.security.checkState()==Qt.CheckState.Checked:
                                    securityThing = "حكومي"
                                elif self.private.checkState()==Qt.CheckState.Checked:
                                    securityThing = "خاص"

                                cr.execute(f"INSERT INTO vactions (employeeId,vactionType,day,execute,Fares,notes,whereFrom,whereExact,HijryDate,GeoDate) VALUES ('{self.identyEntry.text()}','{self.vications.currentText()}','{self.dayEntry.text()}','{ExecuseThing}','{FaresThing}','{self.NotesEntry.text()}','{securityThing}','{self.WhereEntry.text()}','{self.dateHigryEntry.text()}','{self.dateEntry.text()}')")
                                con.commit()          
                                d = QMessageBox(parent=self,text="تم الاضافة بنجاح")
                                d.setWindowTitle("نجاح")
                                d.setIcon(QMessageBox.Icon.Information)
                                d.exec()
                            
                        elif self.otherVaction:
                            if len(self.typeEntry.text()) > 0:
                                cr.execute(f"INSERT INTO vactions (employeeId,vactionType,day,execute,Fares,notes,whereFrom,whereExact,HijryDate,GeoDate) VALUES ('{self.identyEntry.text()}','{self.typeEntry.text()}','{self.dayEntry.text()}','{ExecuseThing}','','{self.NotesEntry.text()}','','','{self.dateHigryEntry.text()}','{self.dateEntry.text()}')")
                                con.commit()          
                                d = QMessageBox(parent=self,text="تم الاضافة بنجاح")
                                d.setWindowTitle("نجاح")
                                d.setIcon(QMessageBox.Icon.Information)
                                d.exec()
                        else:
                            cr.execute(f"INSERT INTO vactions (employeeId,vactionType,day,execute,Fares,notes,whereFrom,whereExact,HijryDate,GeoDate) VALUES ('{self.identyEntry.text()}','{self.vications.currentText()}','{self.dayEntry.text()}','{ExecuseThing}','','{self.NotesEntry.text()}','','','{self.dateHigryEntry.text()}','{self.dateEntry.text()}')")
                            con.commit()          
                            d = QMessageBox(parent=self,text="تم الاضافة بنجاح")
                            d.setWindowTitle("نجاح")
                            d.setIcon(QMessageBox.Icon.Information)
                            d.exec()
                else:
                    if self.withExecuse.checkState()==Qt.CheckState.Checked:
                        ExecuseThing = "بعذر"
                    elif self.withOutExecuse.checkState()==Qt.CheckState.Checked:
                        ExecuseThing = "بدون بعذر"
                    if self.addtional:
                        if (self.withFares.checkState()==Qt.CheckState.Checked or self.withOutFares.checkState()==Qt.CheckState.Checked) and (self.security.checkState()==Qt.CheckState.Checked or self.private.checkState()==Qt.CheckState.Checked) and (len(self.WhereEntry.text()) > 0):
                            if self.withFares.checkState()==Qt.CheckState.Checked:
                                FaresThing = "تم"
                            elif self.withOutFares.checkState()==Qt.CheckState.Checked:
                                FaresThing = "لم يتم"

                            if self.security.checkState()==Qt.CheckState.Checked:
                                securityThing = "حكومي"
                            elif self.private.checkState()==Qt.CheckState.Checked:
                                securityThing = "خاص"
                            cr.execute(f"INSERT INTO vactions (employeeId,vactionType,day,execute,Fares,notes,whereFrom,whereExact,HijryDate,GeoDate) VALUES ('{self.identyEntry.text()}','{self.vications.currentText()}','{self.dayEntry.text()}','{ExecuseThing}','{FaresThing}','{self.NotesEntry.text()}','{securityThing}','{self.WhereEntry.text()}','{self.dateHigryEntry.text()}','{self.dateEntry.text()}')")
                            con.commit()          
                            d = QMessageBox(parent=self,text="تم الاضافة بنجاح")
                            d.setWindowTitle("نجاح")
                            d.setIcon(QMessageBox.Icon.Information)
                            d.exec()
                    elif self.otherVaction:
                        if len(self.typeEntry.text()) > 0:
                            cr.execute(f"INSERT INTO vactions (employeeId,vactionType,day,execute,Fares,notes,whereFrom,whereExact,HijryDate,GeoDate) VALUES ('{self.identyEntry.text()}','{self.typeEntry.text()}','{self.dayEntry.text()}','{ExecuseThing}','','{self.NotesEntry.text()}','','','{self.dateHigryEntry.text()}','{self.dateEntry.text()}')")
                            con.commit()          
                            d = QMessageBox(parent=self,text="تم الاضافة بنجاح")
                            d.setWindowTitle("نجاح")
                            d.setIcon(QMessageBox.Icon.Information)
                            d.exec()
                    else:
                        cr.execute(f"INSERT INTO vactions (employeeId,vactionType,day,execute,Fares,notes,whereFrom,whereExact,HijryDate,GeoDate) VALUES ('{self.identyEntry.text()}','{self.vications.currentText()}','{self.dayEntry.text()}','{ExecuseThing}','','{self.NotesEntry.text()}','','','{self.dateHigryEntry.text()}','{self.dateEntry.text()}')")
                        con.commit()          
                        d = QMessageBox(parent=self,text="تم الاضافة بنجاح")
                        d.setWindowTitle("نجاح")
                        d.setIcon(QMessageBox.Icon.Information)
                        d.exec()
        def Change(self,num):
            if num==0:
                self.withOutExecuse.setCheckState(Qt.CheckState.Unchecked)
            elif num==1:
                self.withExecuse.setCheckState(Qt.CheckState.Unchecked)

            if num==2:
                self.withOutFares.setCheckState(Qt.CheckState.Unchecked)
            elif num==3:
                self.withFares.setCheckState(Qt.CheckState.Unchecked)

            if num==4:
                self.private.setCheckState(Qt.CheckState.Unchecked)
            elif num==5:
                self.security.setCheckState(Qt.CheckState.Unchecked)
        def addtionalThings(self):
            if self.vications.currentText() == "أخرى":
                try:
                    self.AddtionalThingsFrame.destroy()
                    self.AddtionalThingsFrame.close()
                except:
                    pass

                self.otherVaction = True
                self.addtional = False

                self.AddtionalThingsFrame = QFrame(self.frame)
                self.AddtionalThingsFrame.setGeometry(0,440,291,90)
                self.AddtionalThingsFrame.setStyleSheet("background-color:transparent")
                typeLabel = QLabel("الاجازة",self.AddtionalThingsFrame)
                typeLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
                typeLabel.setFont(QFont("Arial",20))
                typeLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
                typeLabel.setGeometry(10,0,270,25)

                self.typeEntry = QLineEdit(self.AddtionalThingsFrame)
                self.typeEntry.setFont(QFont("Arial",18))
                self.typeEntry.setGeometry(10,35,270,25)
                self.typeEntry.setStyleSheet("background-color:white")
                self.AddtionalThingsFrame.show()
            else:
                try:
                    self.AddtionalThingsFrame.destroy()
                    self.AddtionalThingsFrame.close()
                except:
                    pass
                self.addtional = False
                self.otherVaction = False
        def addIssueEmployee(self):
            for i in self.frame.children():
                i.deleteLater()

            self.frame.setGeometry(self.frame.x(),130,291,550)
            self.searchEntry.setDisabled(True)
            self.showRadioButtonsToEdit("Appcent")

            self.imagesPath = []
            identyLabel = QLabel("السجل المدني",self.frame)
            identyLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            identyLabel.setFont(QFont("Arial",20))
            identyLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            identyLabel.setGeometry(10,0,270,30)
            
            self.identyEntry = QLineEdit(self.frame)
            self.identyEntry.setFont(QFont("Arial",18))
            self.identyEntry.setGeometry(10,35,270,30)
            self.identyEntry.setStyleSheet("background-color:white")
            self.identyEntry.setDisabled(True)

            dateLabel = QLabel("التاريخ الميلادي",self.frame)
            dateLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            dateLabel.setFont(QFont("Arial",20))
            dateLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            dateLabel.setGeometry(10,70,270,30)


            self.dateEntry = QLineEdit(self.frame)
            self.dateEntry.setFont(QFont("Arial",18))
            self.dateEntry.setGeometry(10,105,270,30)
            self.dateEntry.setStyleSheet("background-color:white")
            self.dateEntry.setText(CURRENT_TIME_GEO)
            self.dateEntry.setDisabled(True)

            dateHigryLabel = QLabel("التاريخ الهجري",self.frame)
            dateHigryLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            dateHigryLabel.setFont(QFont("Arial",20))
            dateHigryLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            dateHigryLabel.setGeometry(10,140,270,30)

            self.dateHigryEntry = QLineEdit(self.frame)
            self.dateHigryEntry.setFont(QFont("Arial",18))
            self.dateHigryEntry.setGeometry(10,175,270,30)
            self.dateHigryEntry.setStyleSheet("background-color:white")
            self.dateHigryEntry.setText(CURRENT_TIME_HIJ)
            self.dateHigryEntry.setDisabled(True)

            ImageLabel = QLabel("صورة المسائلة",self.frame)
            ImageLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            ImageLabel.setFont(QFont("Arial",20))
            ImageLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            ImageLabel.setGeometry(10,210,270,30)

            self.frmaeImageToIssye = QFrame(self.frame)
            self.frmaeImageToIssye.setGeometry(20,245,250,200)
            self.frmaeImageToIssye.setStyleSheet("background-color:white")
            self.frmaeImageToIssyeLayout = QVBoxLayout(self.frame)
            self.frmaeImageToIssye.setLayout(self.frmaeImageToIssyeLayout)

            ButtonAddPic = QPushButton(self.frame,clicked=self.addPictoAddIsuue)
            ButtonAddPic.setStyleSheet("qproperty-icon:url(assests/cam.png);qproperty-iconSize:30px 30px")
            ButtonAddPic.setGeometry(100,445,30,30)

            deletePicButton = QPushButton(self.frame,clicked=self.deletePictoAddIsuue)
            deletePicButton.setStyleSheet("qproperty-icon:url(assests/trashIcon.png);qproperty-iconSize:30px 30px")
            deletePicButton.setGeometry(170,445,30,30)

            completeAddIssueButton = QPushButton("اضافة",self.frame)
            completeAddIssueButton.setFont(QFont("Arial",20))
            completeAddIssueButton.setGeometry(10,480,270,30)
            completeAddIssueButton.setStyleSheet("background-color:#389D63;color:white")

            completeAddIssueButton.clicked.connect(self.completeAddIssue)

            backButton = QPushButton("العودة",self.frame,clicked=self.showComponents)
            backButton.setFont(QFont("Arial",20))
            backButton.setGeometry(15,515,260,30)
            backButton.setStyleSheet("background-color:#389D63;color:white")

            identyLabel.show()
            self.identyEntry.show()
            dateLabel.show()
            self.dateEntry.show()
            dateHigryLabel.show()
            self.dateHigryEntry.show()
            ImageLabel.show()
            self.frmaeImageToIssye.show()
            ButtonAddPic.show()
            deletePicButton.show()
            completeAddIssueButton.show()
            backButton.show()
            self.frame.show()
        def addPictoAddIsuue(self):
            ImagePath = QFileDialog.getOpenFileName(parent=self,caption="Select a File",filter="Image File (*.png *.jpg *.jpeg)")
            if len(ImagePath[0]) > 0:
                for i in reversed(range(self.frmaeImageToIssyeLayout.count())): 
                        self.frmaeImageToIssyeLayout.itemAt(i).widget().setParent(None)
                        self.imagesPath = []
                image = Image.open(ImagePath[0])
                tempImage = QPixmap(ImagePath[0])
                imageLabel = QLabel(self.frmaeImageToIssye)
                imageLabel.setFixedSize(230,180)
                imageLabel.setScaledContents(True)
                imageLabel.setPixmap(tempImage)
                self.imagesPath.append(ImagePath[0])

                self.frmaeImageToIssyeLayout.addWidget(imageLabel)
        def deletePictoAddIsuue(self):
            for i in reversed(range(self.frmaeImageToIssyeLayout.count())): 
                self.frmaeImageToIssyeLayout.itemAt(i).widget().setParent(None)
                self.imagesPath = []

        def completeAddIssue(self):
            if len(self.identyEntry.text()) > 0 and len(self.dateEntry.text()) > 0 and len(self.dateHigryEntry.text()) > 0 and self.frmaeImageToIssyeLayout.count() > 0:
                with open(self.imagesPath[0],"rb") as binary_image:
                    binaryCode = binary_image.read()
                typeFile = os.path.splitext(self.imagesPath[0])[1]
                cr.execute("SELECT seq FROM sqlite_sequence WHERE name = 'issues'")
                cr.execute("""INSERT INTO imagesIssues (image,issuesId,state) values(?,?,?)""",(binaryCode,str((int(cr.fetchone()[0])+1)),typeFile))
                cr.execute(f"SELECT name FROM info WHERE identy='{self.identyEntry.text()}'")
                cr.execute(f"INSERT INTO issues (employeeIdenty,employeeName,HijryDate,GeoDate) values ('{self.identyEntry.text()}','{cr.fetchone()[0]}','{self.dateHigryEntry.text()}','{self.dateEntry.text()}')")
                con.commit()
                d = QMessageBox(parent=self,text="تم الاضافة بنجاح")
                d.setIcon(QMessageBox.Icon.Information)
                d.setWindowTitle("نجاح")
                d.exec()
                self.showComponents()

        def exportOne(self):
            # self.frame.setGeometry(820,340,291,561) 820,120,291,590
            self.frame.setGeometry(self.frame.x(),190,291,480)
            self.showRadioButtonsToEdit("Appcent")

            for i in self.frame.children():
                i.deleteLater()

            identyLabel = QLabel("السجل المدني",self.frame)
            identyLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            identyLabel.setFont(QFont("Arial",18))
            identyLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            identyLabel.setGeometry(10,0,270,30)

            self.identyEntry = QLineEdit(self.frame)
            self.identyEntry.setFont(QFont("Arial",20))
            self.identyEntry.setGeometry(10,35,270,30)
            self.identyEntry.setStyleSheet("background-color:white")

            beginDateLabel = QLabel("تاريخ البداية",self.frame)
            beginDateLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            beginDateLabel.setFont(QFont("Arial",18))
            beginDateLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            beginDateLabel.setGeometry(10,70,270,30)

            warningLabel1 = QLabel("يجب كتابة التاريخ بهذا التنسيق اليوم/الشهر/السنة",self.frame)
            warningLabel1.setStyleSheet("background:transparent;width:250px;color:red")
            warningLabel1.setFont(QFont("Arial",10))
            warningLabel1.setAlignment(Qt.AlignmentFlag.AlignCenter)
            warningLabel1.setGeometry(10,105,270,20)

            self.beginDateEntry = QLineEdit(self.frame)
            self.beginDateEntry.setFont(QFont("Arial",20))
            self.beginDateEntry.setGeometry(10,130,270,30)
            self.beginDateEntry.setStyleSheet("background-color:white")

            endDateLabel = QLabel("تاريخ النهاية",self.frame)
            endDateLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            endDateLabel.setFont(QFont("Arial",18))
            endDateLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            endDateLabel.setGeometry(10,165,270,30)

            warningLabel = QLabel("يجب كتابة التاريخ بهذا التنسيق اليوم/الشهر/السنة",self.frame)
            warningLabel.setStyleSheet("background:transparent;width:250px;color:red")
            warningLabel.setFont(QFont("Arial",10))
            warningLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            warningLabel.setGeometry(10,190,270,20)

            self.endDateEntry = QLineEdit(self.frame)
            self.endDateEntry.setFont(QFont("Arial",20))
            self.endDateEntry.setGeometry(10,215,270,30)
            self.endDateEntry.setStyleSheet("background-color:white")

            extensiosLabel = QLabel("اختر الصيغة",self.frame)
            extensiosLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            extensiosLabel.setFont(QFont("Arial",18))
            extensiosLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            extensiosLabel.setGeometry(10,250,270,30)

            self.vicationsExten = QComboBox(self.frame)
            self.vicationsExten.setFont(QFont("Arial",15))
            self.vicationsExten.setGeometry(10,285,270,30)
            self.vicationsExten.setStyleSheet("background-color:white")

            self.vicationsExten.addItems(["Word","Pdf"])


            completeExportOne = QPushButton("تصدير",self.frame)
            completeExportOne.setFont(QFont("Arial",20))
            completeExportOne.setGeometry(10,320,270,40)
            completeExportOne.setStyleSheet("background-color:#389D63;color:white")
            completeExportOne.clicked.connect(self.completeExportOneFun)

            backButton = QPushButton("العودة",self.frame)
            backButton.setFont(QFont("Arial",20))
            backButton.setGeometry(10,370,270,40)
            backButton.setStyleSheet("background-color:#389D63;color:white")
            backButton.clicked.connect(self.showComponents)

            identyLabel.show()
            self.identyEntry.show()
            beginDateLabel.show()
            warningLabel.show()
            warningLabel1.show()
            self.beginDateEntry.show()
            endDateLabel.show()
            warningLabel.show()
            self.endDateEntry.show()
            extensiosLabel.show()
            self.vicationsExten.show()
            completeExportOne.show()
            backButton.show()
            self.frame.show()
        def completeExportOneFun(self):
            if len(self.identyEntry.text()) > 0 and len(self.beginDateEntry.text()) > 0 and len(self.endDateEntry.text()) > 0 and len (self.vicationsExten.currentText()) > 0:
                try:
                    error = False
                    startDateBeginThing = (str(self.beginDateEntry.text())).split("/")
                    finishDateHijryReady = Hijri(int(startDateBeginThing[0]),int(startDateBeginThing[1]),int(startDateBeginThing[2]))

                    startDateEndThing = (str(self.endDateEntry.text())).split("/")
                    finishEndDateHijryReady = Hijri(int(startDateEndThing[0]),int(startDateEndThing[1]),int(startDateEndThing[2]))
                except:
                    error = True
                    d = QMessageBox(parent=self,text="هناك خطأ في تنسيق التاريخ")
                    d.setWindowTitle("ERROR")
                    d.setIcon(QMessageBox.Icon.Critical)
                    d.exec()
                if not error:
                    try:
                        Dates = []
                        cr.execute(f"SELECT HijryDate From vactions WHERE employeeId='{self.identyEntry.text()}'")

                        for i in cr.fetchall():
                            for j in i:
                                startDateBeginThingTemp = (str(j)).split("/")
                                finishDateHijryReadyTemp = Hijri(int(startDateBeginThingTemp[0]),int(startDateBeginThingTemp[1]),int(startDateBeginThingTemp[2]))

                                if finishDateHijryReadyTemp >=finishDateHijryReady and finishDateHijryReadyTemp<=finishEndDateHijryReady:
                                    if j not in Dates:
                                        Dates.append(j)

                        if self.vicationsExten.currentText()=="Word":
                            self.writeWord(Dates)
                        else:
                            self.writeWord(Dates,"Pdf")
                    except:
                        pass    
        def writeWord(self,dates,what="Word",count="Single"):
            if count=="Single":
                cr.execute(f"SELECT name FROM info WHERE identy='{self.identyEntry.text()}'")
                nameSpicailThingImportant = cr.fetchone()[0]
            file_path = QFileDialog.getExistingDirectory(self,"Select a Directory")
            if len(file_path) > 0:
                doc = docx.Document()
                sections = doc.sections
                sections.page_height = 11.69
                sections.page_width = 8.27
                sections = sections[-1]
                sections.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE


                new_width,new_height = sections.page_height,sections.page_width
                sections.page_width = new_width
                sections.page_height = new_height

                sections = doc.sections

                for section in sections:
                    section.top_margin = docx.shared.Cm(0.3)
                    section.bottom_margin = docx.shared.Cm(0.3)
                    section.left_margin = docx.shared.Cm(0.3)
                    section.right_margin = docx.shared.Cm(0.3)
                

                employees_table = doc.add_table(rows=1,cols=9)
                employees_table.style = "Table Grid"
                hdr_Cells = employees_table.rows[0].cells
                hdr_Cells[8].text = "م"
                hdr_Cells[7].text = "اسم المتطوع"
                hdr_Cells[6].text = "نوع الاجازة"
                hdr_Cells[5].text = "اليوم"
                hdr_Cells[4].text = "ملاحظات"
                hdr_Cells[3].text = "نوع التقرير"
                hdr_Cells[2].text = "التاريخ الهجري"
                hdr_Cells[1].text = "التاريخ الميلادي"
                hdr_Cells[0].text = "الصورة"

                
                b = 0
                tempPictures = []
                for i in dates:
                    if count=="Single":
                        cr.execute(f"SELECT employeeId,vactionType,day,execute,notes,HijryDate,GeoDate,id FROM vactions WHERE HijryDate = '{i}' and employeeId='{self.identyEntry.text()}'")
                    else:
                        cr.execute(f"SELECT employeeId,vactionType,day,execute,notes,HijryDate,GeoDate,id FROM vactions WHERE HijryDate = '{i}'")
                    for jx in cr.fetchall():
                        b+=1   
                        row_Cells = employees_table.add_row().cells
                        row_Cells[0].size = docx.shared.Pt(15)
                        row_Cells[1].size = docx.shared.Pt(15)
                        row_Cells[2].size = docx.shared.Pt(15)
                        row_Cells[3].size = docx.shared.Pt(15)
                        row_Cells[4].size = docx.shared.Pt(15)
                        row_Cells[5].size = docx.shared.Pt(15)
                        row_Cells[6].size = docx.shared.Pt(15)
                        row_Cells[7].size = docx.shared.Pt(15)
                        row_Cells[8].size = docx.shared.Pt(15)

                        row_Cells[0].rtl = True
                        row_Cells[1].rtl = True
                        row_Cells[2].rtl = True
                        row_Cells[3].rtl = True
                        row_Cells[4].rtl = True
                        row_Cells[5].rtl = True
                        row_Cells[6].rtl = True
                        row_Cells[7].rtl = True
                        row_Cells[8].rtl = True



                        cr.execute(f"SELECT name FROM info WHERE identy='{jx[0]}'")
                        tempNameToExport = cr.fetchone()[0]
                        row_Cells[8].text = str(b)
                        row_Cells[7].text = str(tempNameToExport)
                        row_Cells[6].text = str(jx[1])
                        row_Cells[5].text = str(jx[2])
                        row_Cells[4].text = str(jx[4])
                        row_Cells[3].text = str(jx[3])
                        row_Cells[2].text = str(jx[5])
                        row_Cells[1].text = str(jx[6])


                        if row_Cells[1].text == "":
                            cell_xml_element = row_Cells[1]._tc
                            table_cell_properties = cell_xml_element.get_or_add_tcPr()
                            shade_obj = OxmlElement("w:shd")
                            shade_obj.set(qn2("w:fill"),"F3FF00")
                            table_cell_properties.append(shade_obj)
                        if row_Cells[2].text == "":
                            cell_xml_element = row_Cells[2]._tc
                            table_cell_properties = cell_xml_element.get_or_add_tcPr()
                            shade_obj = OxmlElement("w:shd")
                            shade_obj.set(qn2("w:fill"),"F3FF00")
                            table_cell_properties.append(shade_obj)
                        if row_Cells[3].text == "":
                            cell_xml_element = row_Cells[3]._tc
                            table_cell_properties = cell_xml_element.get_or_add_tcPr()
                            shade_obj = OxmlElement("w:shd")
                            shade_obj.set(qn2("w:fill"),"F3FF00")
                            table_cell_properties.append(shade_obj)
                        if row_Cells[4].text == "":
                            cell_xml_element = row_Cells[4]._tc
                            table_cell_properties = cell_xml_element.get_or_add_tcPr()
                            shade_obj = OxmlElement("w:shd")
                            shade_obj.set(qn2("w:fill"),"F3FF00")
                            table_cell_properties.append(shade_obj)
                        if row_Cells[5].text == "":
                            cell_xml_element = row_Cells[5]._tc
                            table_cell_properties = cell_xml_element.get_or_add_tcPr()
                            shade_obj = OxmlElement("w:shd")
                            shade_obj.set(qn2("w:fill"),"F3FF00")
                            table_cell_properties.append(shade_obj)
                        if row_Cells[6].text == "":
                            cell_xml_element = row_Cells[6]._tc
                            table_cell_properties = cell_xml_element.get_or_add_tcPr()
                            shade_obj = OxmlElement("w:shd")
                            shade_obj.set(qn2("w:fill"),"F3FF00")
                            table_cell_properties.append(shade_obj)
                        if row_Cells[7].text == "":
                            cell_xml_element = row_Cells[7]._tc
                            table_cell_properties = cell_xml_element.get_or_add_tcPr()
                            shade_obj = OxmlElement("w:shd")
                            shade_obj.set(qn2("w:fill"),"F3FF00")
                            table_cell_properties.append(shade_obj)
                        if row_Cells[8].text == "":
                            cell_xml_element = row_Cells[8]._tc
                            table_cell_properties = cell_xml_element.get_or_add_tcPr()
                            shade_obj = OxmlElement("w:shd")
                            shade_obj.set(qn2("w:fill"),"F3FF00")
                            table_cell_properties.append(shade_obj)


                        cr.execute(f"SELECT * FROM images where vactionId='{jx[7]}'")
                        tempValuesImages = cr.fetchone()

                        if tempValuesImages !=None:
                            with open(f"tempimages/{tempValuesImages[1]}{tempValuesImages[2]}","wb") as image:
                                image.write(tempValuesImages[0])
                            paragraph =hdr_Cells[0].paragraphs[0]
                            run = paragraph.runs
                            font = run[0].font
                            font.size= docx.shared.Pt(15)
                            cells = row_Cells[0].paragraphs[0]
                            cells.add_run().add_picture(f"tempimages/{tempValuesImages[1]}{tempValuesImages[2]}",width=docx.shared.Inches(1),height=docx.shared.Inches(1))
                        else:
                            cell_xml_element = row_Cells[0]._tc
                            table_cell_properties = cell_xml_element.get_or_add_tcPr()
                            shade_obj = OxmlElement("w:shd")
                            shade_obj.set(qn2("w:fill"),"F3FF00")
                            table_cell_properties.append(shade_obj)

                for row in employees_table.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                widths = (docx.shared.Inches(4),docx.shared.Inches(3),docx.shared.Inches(3),docx.shared.Inches(2),docx.shared.Inches(3),docx.shared.Inches(2),docx.shared.Inches(3), docx.shared.Inches(4),docx.shared.Inches(0.5))
                for row in employees_table.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width
                

                for row in employees_table.rows:
                    for cell in row.cells:
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(17)
                                
                if count=="Single":
                    doc.save(f"{file_path}/معلومات الغياب ل {nameSpicailThingImportant}.docx")
                else:
                    doc.save(f"{file_path}/معلومات الغياب.docx")

                if what=="Pdf":
                    if count=="Single":
                        
                        in_file = str(os.path.abspath(f"{file_path}/معلومات الغياب ل {nameSpicailThingImportant}.docx")).replace("c","C")
                        out_file = str(os.path.abspath(f"{file_path}/معلومات الغياب ل {nameSpicailThingImportant}.pdf")).replace("c","C")

                        TempdocO = aw.Document(in_file)
                        TempdocO.save(out_file)

                    else:
                        in_file = str(os.path.abspath(f"{file_path}/معلومات الغياب.docx")).replace("c","C")
                        out_file = str(os.path.abspath(f"{file_path}/معلومات الغياب.pdf")).replace("c","C")

                        TempdocO = aw.Document(in_file)
                        TempdocO.save(out_file)

                    try:
                        if count=="Single":
                            os.remove(f"{file_path}/معلومات الغياب ل {nameSpicailThingImportant}.docx")
                        else:
                            os.remove(f"{file_path}/معلومات الغياب.docx")
                    except:
                        pass

                for i in os.listdir("tempImages"):
                    try:
                        os.remove(f"tempImages/{i}")
                    except:
                        pass
                d = QMessageBox(parent=self,text="تم الحفظ بنجاح")
                d.setWindowTitle("نجاح")
                d.setIcon(QMessageBox.Icon.Information)
                d.exec()
        def exportAll(self):
            for i in self.frame.children():
                i.deleteLater()
            self.frame.setGeometry(self.frame.x(),190,291,480)

            beginDateLabel = QLabel("تاريخ البداية",self.frame)
            beginDateLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            beginDateLabel.setFont(QFont("Arial",18))
            beginDateLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            beginDateLabel.setGeometry(10,0,270,30)

            warningLabel1 = QLabel("يجب كتابة التاريخ بهذا التنسيق اليوم/الشهر/السنة",self.frame)
            warningLabel1.setStyleSheet("background:transparent;width:250px;color:red")
            warningLabel1.setFont(QFont("Arial",10))
            warningLabel1.setAlignment(Qt.AlignmentFlag.AlignCenter)
            warningLabel1.setGeometry(10,35,270,20)

            self.beginDateEntry = QLineEdit(self.frame)
            self.beginDateEntry.setFont(QFont("Arial",20))
            self.beginDateEntry.setGeometry(10,70,270,30)
            self.beginDateEntry.setStyleSheet("background-color:white")

            endDateLabel = QLabel("تاريخ النهاية",self.frame)
            endDateLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            endDateLabel.setFont(QFont("Arial",18))
            endDateLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            endDateLabel.setGeometry(10,105,270,30)

            warningLabel = QLabel("يجب كتابة التاريخ بهذا التنسيق اليوم/الشهر/السنة",self.frame)
            warningLabel.setStyleSheet("background:transparent;width:250px;color:red")
            warningLabel.setFont(QFont("Arial",10))
            warningLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            warningLabel.setGeometry(10,140,270,20)

            self.endDateEntry = QLineEdit(self.frame)
            self.endDateEntry.setFont(QFont("Arial",20))
            self.endDateEntry.setGeometry(10,175,270,30)
            self.endDateEntry.setStyleSheet("background-color:white")

            extensiosLabel = QLabel("اختر الصيغة",self.frame)
            extensiosLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            extensiosLabel.setFont(QFont("Arial",18))
            extensiosLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            extensiosLabel.setGeometry(10,210,270,30)

            self.vicationsExten = QComboBox(self.frame)
            self.vicationsExten.setFont(QFont("Arial",15))
            self.vicationsExten.setGeometry(10,245,270,30)
            self.vicationsExten.setStyleSheet("background-color:white")

            self.vicationsExten.addItems(["Word","Pdf"])


            completeExportOne = QPushButton("تصدير",self.frame)
            completeExportOne.setFont(QFont("Arial",20))
            completeExportOne.setGeometry(10,280,270,40)
            completeExportOne.setStyleSheet("background-color:#389D63;color:white")
            completeExportOne.clicked.connect(self.completeexportAll)

            backButton = QPushButton("العودة",self.frame)
            backButton.setFont(QFont("Arial",20))
            backButton.setGeometry(10,330,270,40)
            backButton.setStyleSheet("background-color:#389D63;color:white")
            backButton.clicked.connect(self.showComponents)

            beginDateLabel.show()
            warningLabel1.show()
            self.beginDateEntry.show()
            endDateLabel.show()
            warningLabel.show()
            self.endDateEntry.show()
            extensiosLabel.show()
            self.vicationsExten.show()
            completeExportOne.show()
            backButton.show()
            self.frame.show()
        def completeexportAll(self):
            if len(self.beginDateEntry.text()) > 0 and len(self.endDateEntry.text()) > 0 and len (self.vicationsExten.currentText()) > 0:
                try:
                    error = False
                    startDateBeginThing = (str(self.beginDateEntry.text())).split("/")
                    finishDateHijryReady = Hijri(int(startDateBeginThing[0]),int(startDateBeginThing[1]),int(startDateBeginThing[2]))

                    startDateEndThing = (str(self.endDateEntry.text())).split("/")
                    finishEndDateHijryReady = Hijri(int(startDateEndThing[0]),int(startDateEndThing[1]),int(startDateEndThing[2]))
                except:
                    error = True
                    d = QMessageBox(parent=self,text="هناك خطأ في تنسيق التاريخ")
                    d.setWindowTitle("ERROR")
                    d.setIcon(QMessageBox.Icon.Critical)
                    d.exec()
                if not error:
                    # try:
                        Dates = []
                        cr.execute(f"SELECT HijryDate From vactions")


                        for i in cr.fetchall():
                            for j in i:
                                startDateBeginThingTemp = (str(j)).split("/")
                                finishDateHijryReadyTemp = Hijri(int(startDateBeginThingTemp[0]),int(startDateBeginThingTemp[1]),int(startDateBeginThingTemp[2]))

                                if finishDateHijryReadyTemp >=finishDateHijryReady and finishDateHijryReadyTemp<=finishEndDateHijryReady:
                                    if j not in Dates:
                                        Dates.append(j)

                        if self.vicationsExten.currentText()=="Word":
                            self.writeWord(Dates,"Word","All")
                        else:
                            self.writeWord(Dates,"Pdf","All")
                    # except:
                    #     pass   

        def insertDatabaseFun(self):
            self.confirmWindow = tempwindows()
            self.confirmWindow.setFixedSize(394,200)
            message = "عند استيراد قاعدة بيانات                                 \nسوف تنحذف جميع المعلومات من قاعدة البيانات ولايمكن استرجاعها\n                            يرجى كتابة 'تأكيد'بالحقل للتنفيذ"
            problems = QLabel(message,self.confirmWindow)
            problems.setStyleSheet("color:red;font-size:14px")

            self.confirmEntry = QLineEdit(self.confirmWindow)
            self.confirmEntry.setFont(QFont("Arial",15))
            self.confirmEntry.setGeometry(103,80,200,30)

            confirmbutton = QPushButton("تنفيذ",self.confirmWindow,clicked=self.completeinsertDatabaseFun)
            confirmbutton.setFont(QFont("Arial",20))
            confirmbutton.setStyleSheet("background-color:#389D63;color:white")
            confirmbutton.setGeometry(103,130,200,35)

            self.confirmWindow.show()
        def completeinsertDatabaseFun(self):
            if self.confirmEntry.text() == "تأكيد":
                fileDbUser = QFileDialog.getOpenFileName(parent=self.confirmWindow,caption="Select a File",filter="Database File (*.db)")
                if len(fileDbUser[0]) > 0:
                    try:
                        """
                        #con1 = sqlite3.connect(fileDbUser[0])
                        cr1 = con1.cursor()
                        cr1.execute("SELECT useAble FROM confirmationDatabase")
                        if cr1.fetchone()[0] == "canUse":
                        it's suposed to be like this
                        """
                        cr.execute("SELECT useAble FROM confirmationDatabase")
                        if cr.fetchone()[0] == "canUse":
                            con.close()
                            os.remove("app.db")
                            shutil.copy2(fileDbUser[0],f"app.db")
                            d = QMessageBox(parent=self,text="تم الحفظ بنجاح")
                            d.setWindowTitle("نجاح")
                            d.setIcon(QMessageBox.Icon.Information)
                            d.exec()
                            self.confirmWindow.destroy()
                            self.destroy()
                            mainToRunApp()
                        else:
                            raise Exception("notUseAble")
                    except:
                        d = QMessageBox(parent=self,text="قاعدة البيانات غير صالحة")
                        d.setWindowTitle("ERROR")
                        d.setIcon(QMessageBox.Icon.Critical)
                        d.exec()
        def exportDatabaseFun(self):
            filePath = QFileDialog.getExistingDirectory(self,"Select a Directory")
            if len(filePath)> 0:
                shutil.copy2("app.db",f"{filePath}")
                d = QMessageBox(parent=self,text="تم الحفظ بنجاح")
                d.setWindowTitle("نجاح")
                d.setIcon(QMessageBox.Icon.Information)
                d.exec()
        def issuesDetalis(self):
            self.delatedItems = []
            
            self.appcentsDetalisWindow = tempwindows()
            self.appcentsDetalisWindow.setFixedSize(500,500)
            self.appcentsDetalisTable = QTableWidget(self.appcentsDetalisWindow)
            self.appcentsDetalisTable.setGeometry(10,40,456,400)
            self.appcentsDetalisTable.setColumnCount(6) 
            self.appcentsDetalisTable.setColumnHidden(5,True)
            self.appcentsDetalisTable.setColumnWidth(0,30)
            self.appcentsDetalisTable.setHorizontalHeaderLabels(["","اسم المتطوع","التاريخ الهجري","التاريخ الميلادي","الصورة",""])
            
            buttonSaveEveryThing = QPushButton(self.appcentsDetalisWindow,clicked=self.comleteEveryThingINissuesTable)
            buttonSaveEveryThing.setStyleSheet("background-color:green;color:white")
            buttonSaveEveryThing.setText("حفظ التغيرات")
            buttonSaveEveryThing.setGeometry(160,450,150,40)

            cr.execute(f"SELECT * FROM issues WHERE employeeIdenty='{self.reportEntry.text()}'")
            infos = []
            for i in cr.fetchall():
                listi = list(i)
                cr.execute(f"SELECT name FROM info WHERE identy='{listi[0]}'")
                listi[0] = cr.fetchone()[0]
                listi[-1] = str(listi[-1])
                listi.insert(0,"")
                infos.append(listi)
            for row,i in enumerate(infos):
                self.appcentsDetalisTable.insertRow(self.appcentsDetalisTable.rowCount())
                for col,val in enumerate(i):
                    if col%6==0:
                        icon = QPixmap("assests/trashicon.png")
                        button = QPushButton()
                        button.setStyleSheet(f"Qproperty-icon:url(assests/trashicon.png);qproperty-iconSize:30px 30px;background-color:rgb(253, 253, 253)")
                        button.clicked.connect(lambda x,row=row:self.deleteFromIssuesTable(row))
                        self.appcentsDetalisTable.setIndexWidget(self.appcentsDetalisTable.model().index(row,0),button)
                    else:
                        self.appcentsDetalisTable.setItem(row,col,QTableWidgetItem(val))
            for i in range(self.appcentsDetalisTable.rowCount()):
                self.appcentsDetalisTable.item(i,1).setFlags(Qt.ItemFlag.ItemIsEditable)
                self.appcentsDetalisTable.item(i,2).setFlags(Qt.ItemFlag.ItemIsEditable)
                self.appcentsDetalisTable.item(i,3).setFlags(Qt.ItemFlag.ItemIsEditable)
                self.appcentsDetalisTable.item(i,4).setFlags(Qt.ItemFlag.ItemIsEditable)
                self.appcentsDetalisTable.item(i,5).setFlags(Qt.ItemFlag.ItemIsEditable)


            for i in range(self.appcentsDetalisTable.rowCount()):
                cr.execute(f"SELECT image,state FROM imagesIssues where issuesId='{self.appcentsDetalisTable.item(i,5).text()}'")
                imagesData = cr.fetchone()
                if imagesData !=None:
                    item = self.returnQLabelWithImage(imagesData[0],imagesData[1])
                    self.appcentsDetalisTable.setCellWidget(i,4,item)
            self.appcentsDetalisWindow.show()
        def deleteFromIssuesTable(self,row):
            self.delatedItems.append(self.appcentsDetalisTable.item(row,5).text())
            self.appcentsDetalisTable.hideRow(row)
        def comleteEveryThingINissuesTable(self):
            for i in self.delatedItems:
                cr.execute(f"DELETE FROM issues WHERE id='{i}'")
                cr.execute(f"DELETE FROM imagesIssues WHERE issuesId='{i}'")
            con.commit()
        def exportAllIssues(self):
            for i in self.frame.children():
                i.deleteLater()
            self.frame.setGeometry(self.frame.x(),190,291,480)

            beginDateLabel = QLabel("تاريخ البداية",self.frame)
            beginDateLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            beginDateLabel.setFont(QFont("Arial",18))
            beginDateLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            beginDateLabel.setGeometry(10,0,270,30)

            warningLabel1 = QLabel("يجب كتابة التاريخ بهذا التنسيق اليوم/الشهر/السنة",self.frame)
            warningLabel1.setStyleSheet("background:transparent;width:250px;color:red")
            warningLabel1.setFont(QFont("Arial",10))
            warningLabel1.setAlignment(Qt.AlignmentFlag.AlignCenter)
            warningLabel1.setGeometry(10,35,270,20)

            self.beginDateEntry = QLineEdit(self.frame)
            self.beginDateEntry.setFont(QFont("Arial",20))
            self.beginDateEntry.setGeometry(10,70,270,30)
            self.beginDateEntry.setStyleSheet("background-color:white")

            endDateLabel = QLabel("تاريخ النهاية",self.frame)
            endDateLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            endDateLabel.setFont(QFont("Arial",18))
            endDateLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            endDateLabel.setGeometry(10,105,270,30)

            warningLabel = QLabel("يجب كتابة التاريخ بهذا التنسيق اليوم/الشهر/السنة",self.frame)
            warningLabel.setStyleSheet("background:transparent;width:250px;color:red")
            warningLabel.setFont(QFont("Arial",10))
            warningLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            warningLabel.setGeometry(10,140,270,20)

            self.endDateEntry = QLineEdit(self.frame)
            self.endDateEntry.setFont(QFont("Arial",20))
            self.endDateEntry.setGeometry(10,175,270,30)
            self.endDateEntry.setStyleSheet("background-color:white")

            extensiosLabel = QLabel("اختر الصيغة",self.frame)
            extensiosLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            extensiosLabel.setFont(QFont("Arial",18))
            extensiosLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            extensiosLabel.setGeometry(10,210,270,30)

            self.vicationsExten = QComboBox(self.frame)
            self.vicationsExten.setFont(QFont("Arial",15))
            self.vicationsExten.setGeometry(10,245,270,30)
            self.vicationsExten.setStyleSheet("background-color:white")

            self.vicationsExten.addItems(["Word","Pdf"])


            completeExportOne = QPushButton("تصدير",self.frame)
            completeExportOne.setFont(QFont("Arial",20))
            completeExportOne.setGeometry(10,280,270,40)
            completeExportOne.setStyleSheet("background-color:#389D63;color:white")
            completeExportOne.clicked.connect(self.completeExportAllIssues)

            backButton = QPushButton("العودة",self.frame)
            backButton.setFont(QFont("Arial",20))
            backButton.setGeometry(10,330,270,40)
            backButton.setStyleSheet("background-color:#389D63;color:white")
            backButton.clicked.connect(self.showComponents)

            beginDateLabel.show()
            warningLabel1.show()
            self.beginDateEntry.show()
            endDateLabel.show()
            warningLabel.show()
            self.endDateEntry.show()
            extensiosLabel.show()
            self.vicationsExten.show()
            completeExportOne.show()
            backButton.show()
            self.frame.show()
        def completeExportAllIssues(self):
            if len(self.beginDateEntry.text()) > 0 and len(self.endDateEntry.text()) > 0 and len (self.vicationsExten.currentText()) > 0:
                try:    
                    error = False
                    startDateBeginThing = (str(self.beginDateEntry.text())).split("/")
                    finishDateHijryReady = Hijri(int(startDateBeginThing[0]),int(startDateBeginThing[1]),int(startDateBeginThing[2]))

                    startDateEndThing = (str(self.endDateEntry.text())).split("/")
                    finishEndDateHijryReady = Hijri(int(startDateEndThing[0]),int(startDateEndThing[1]),int(startDateEndThing[2]))
                except:
                    error = True
                    d = QMessageBox(parent=self,text="هناك خطأ في تنسيق التاريخ")
                    d.setWindowTitle("ERROR")
                    d.setIcon(QMessageBox.Icon.Critical)
                    d.exec()
                if not error:
                    try:
                        Dates = []
                        cr.execute(f"SELECT HijryDate From issues")


                        for i in cr.fetchall():
                            for j in i:
                                startDateBeginThingTemp = (str(j)).split("/")
                                finishDateHijryReadyTemp = Hijri(int(startDateBeginThingTemp[0]),int(startDateBeginThingTemp[1]),int(startDateBeginThingTemp[2]))

                                if finishDateHijryReadyTemp >=finishDateHijryReady and finishDateHijryReadyTemp<=finishEndDateHijryReady:
                                    if j not in Dates:
                                        Dates.append(j)

                        if self.vicationsExten.currentText()=="Word":
                            self.writeWordIssues(Dates,"Word","All")
                        else:
                            self.writeWordIssues(Dates,"Pdf","All")
                    except:
                        pass   
        def exportSingleIssues(self):
            for i in self.frame.children():
                i.deleteLater()
            self.frame.setGeometry(self.frame.x(),190,291,480)

            self.showRadioButtonsToEdit("Appcent")


            identyLabel = QLabel("السجل المدني",self.frame)
            identyLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            identyLabel.setFont(QFont("Arial",18))
            identyLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            identyLabel.setGeometry(10,0,270,30)

            self.identyEntry = QLineEdit(self.frame)
            self.identyEntry.setFont(QFont("Arial",20))
            self.identyEntry.setGeometry(10,35,270,30)
            self.identyEntry.setStyleSheet("background-color:white")

            beginDateLabel = QLabel("تاريخ البداية",self.frame)
            beginDateLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            beginDateLabel.setFont(QFont("Arial",18))
            beginDateLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            beginDateLabel.setGeometry(10,70,270,30)

            warningLabel1 = QLabel("يجب كتابة التاريخ بهذا التنسيق اليوم/الشهر/السنة",self.frame)
            warningLabel1.setStyleSheet("background:transparent;width:250px;color:red")
            warningLabel1.setFont(QFont("Arial",10))
            warningLabel1.setAlignment(Qt.AlignmentFlag.AlignCenter)
            warningLabel1.setGeometry(10,105,270,20)

            self.beginDateEntry = QLineEdit(self.frame)
            self.beginDateEntry.setFont(QFont("Arial",20))
            self.beginDateEntry.setGeometry(10,130,270,30)
            self.beginDateEntry.setStyleSheet("background-color:white")

            endDateLabel = QLabel("تاريخ النهاية",self.frame)
            endDateLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            endDateLabel.setFont(QFont("Arial",18))
            endDateLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            endDateLabel.setGeometry(10,165,270,30)

            warningLabel = QLabel("يجب كتابة التاريخ بهذا التنسيق اليوم/الشهر/السنة",self.frame)
            warningLabel.setStyleSheet("background:transparent;width:250px;color:red")
            warningLabel.setFont(QFont("Arial",10))
            warningLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            warningLabel.setGeometry(10,190,270,20)

            self.endDateEntry = QLineEdit(self.frame)
            self.endDateEntry.setFont(QFont("Arial",20))
            self.endDateEntry.setGeometry(10,215,270,30)
            self.endDateEntry.setStyleSheet("background-color:white")

            extensiosLabel = QLabel("اختر الصيغة",self.frame)
            extensiosLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            extensiosLabel.setFont(QFont("Arial",18))
            extensiosLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            extensiosLabel.setGeometry(10,250,270,30)

            self.vicationsExten = QComboBox(self.frame)
            self.vicationsExten.setFont(QFont("Arial",15))
            self.vicationsExten.setGeometry(10,285,270,30)
            self.vicationsExten.setStyleSheet("background-color:white")

            self.vicationsExten.addItems(["Word","Pdf"])


            completeExportOne = QPushButton("تصدير",self.frame)
            completeExportOne.setFont(QFont("Arial",20))
            completeExportOne.setGeometry(10,320,270,40)
            completeExportOne.setStyleSheet("background-color:#389D63;color:white")
            completeExportOne.clicked.connect(self.completeExportSingleIssues)

            backButton = QPushButton("العودة",self.frame)
            backButton.setFont(QFont("Arial",20))
            backButton.setGeometry(10,370,270,40)
            backButton.setStyleSheet("background-color:#389D63;color:white")
            backButton.clicked.connect(self.showComponents)

            
            identyLabel.show()
            self.identyEntry.show()
            beginDateLabel.show()
            warningLabel1.show()
            self.beginDateEntry.show()
            endDateLabel.show()
            warningLabel.show()
            self.endDateEntry.show()
            extensiosLabel.show()
            self.vicationsExten.show()
            completeExportOne.show()
            backButton.show()
            self.frame.show()

        def completeExportSingleIssues(self):
            if len(self.identyEntry.text()) > 0 and len(self.beginDateEntry.text()) > 0 and len(self.endDateEntry.text()) > 0 and len (self.vicationsExten.currentText()) > 0:
                try:
                    error = False
                    startDateBeginThing = (str(self.beginDateEntry.text())).split("/")
                    finishDateHijryReady = Hijri(int(startDateBeginThing[0]),int(startDateBeginThing[1]),int(startDateBeginThing[2]))

                    startDateEndThing = (str(self.endDateEntry.text())).split("/")
                    finishEndDateHijryReady = Hijri(int(startDateEndThing[0]),int(startDateEndThing[1]),int(startDateEndThing[2]))
                except:
                    error = True
                    d = QMessageBox(parent=self,text="هناك خطأ في تنسيق التاريخ")
                    d.setWindowTitle("ERROR")
                    d.setIcon(QMessageBox.Icon.Critical)
                    d.exec()
                if not error:
                    try:
                        Dates = []
                        cr.execute(f"SELECT HijryDate From issues WHERE employeeIdenty='{self.identyEntry.text()}'")


                        for i in cr.fetchall():
                            for j in i:
                                startDateBeginThingTemp = (str(j)).split("/")
                                finishDateHijryReadyTemp = Hijri(int(startDateBeginThingTemp[0]),int(startDateBeginThingTemp[1]),int(startDateBeginThingTemp[2]))

                                if finishDateHijryReadyTemp >=finishDateHijryReady and finishDateHijryReadyTemp<=finishEndDateHijryReady:
                                    if j not in Dates:
                                        Dates.append(j)

                        if self.vicationsExten.currentText()=="Word":
                            self.writeWordIssues(Dates)
                        else:
                            self.writeWordIssues(Dates,"Pdf")
                    except:
                        pass
        def writeWordIssues(self,dates,what="Word",count="Single"):
            if count == "Single":
                cr.execute(f"SELECT name FROM info WHERE identy='{self.identyEntry.text()}'")
                nameSpicailThingImportant = cr.fetchone()[0]
            file_path = QFileDialog.getExistingDirectory(self,"Select a Directory")
            if len(file_path) > 0:
                doc = docx.Document()
                sections = doc.sections
                sections.page_height = 11.69
                sections.page_width = 8.27
                sections = sections[-1]
                sections.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE

                new_width,new_height = sections.page_height,sections.page_width
                sections.page_width = new_width
                sections.page_height = new_height

                sections = doc.sections

                for section in sections:
                    section.top_margin = docx.shared.Cm(0.3)
                    section.bottom_margin = docx.shared.Cm(0.3)
                    section.left_margin = docx.shared.Cm(0.3)
                    section.right_margin = docx.shared.Cm(0.3)
                

                employees_table = doc.add_table(rows=1,cols=5)
                employees_table.style = "Table Grid"
                hdr_Cells = employees_table.rows[0].cells
                hdr_Cells[4].text = "م"
                hdr_Cells[3].text = "اسم المتطوع"
                hdr_Cells[2].text = "التاريخ الهجري"
                hdr_Cells[1].text = "التاريخ الميلادي"
                hdr_Cells[0].text = "الصورة"

                
                b = 0
                tempPictures = []
                for i in dates:
                    if count=="Single":
                        cr.execute(f"SELECT employeeName,HijryDate,GeoDate,id FROM issues WHERE HijryDate = '{i}' and employeeIdenty='{self.identyEntry.text()}'")
                    else:
                        cr.execute(f"SELECT employeeName,HijryDate,GeoDate,id FROM issues WHERE HijryDate = '{i}'")
                    for jx in cr.fetchall():
                        b+=1   
                        row_Cells = employees_table.add_row().cells
                        
                        row_Cells[0].size = docx.shared.Pt(15)
                        row_Cells[1].size = docx.shared.Pt(15)
                        row_Cells[2].size = docx.shared.Pt(15)
                        row_Cells[3].size = docx.shared.Pt(15)
                        row_Cells[4].size = docx.shared.Pt(15)

                        row_Cells[4].text = str(b)
                        row_Cells[3].text = str(jx[0])
                        row_Cells[2].text = str(jx[1])
                        row_Cells[1].text = str(jx[2])
                        # row_Cells[0].text = str(jx[3])

                        row_Cells[0].rtl = True
                        row_Cells[1].rtl = True
                        row_Cells[2].rtl = True
                        row_Cells[3].rtl = True
                        row_Cells[4].rtl = True


                        if row_Cells[1].text == "":
                            cell_xml_element = row_Cells[1]._tc
                            table_cell_properties = cell_xml_element.get_or_add_tcPr()
                            shade_obj = OxmlElement("w:shd")
                            shade_obj.set(qn2("w:fill"),"F3FF00")
                            table_cell_properties.append(shade_obj)
                        if row_Cells[2].text == "":
                            cell_xml_element = row_Cells[2]._tc
                            table_cell_properties = cell_xml_element.get_or_add_tcPr()
                            shade_obj = OxmlElement("w:shd")
                            shade_obj.set(qn2("w:fill"),"F3FF00")
                            table_cell_properties.append(shade_obj)
                        if row_Cells[3].text == "":
                            cell_xml_element = row_Cells[3]._tc
                            table_cell_properties = cell_xml_element.get_or_add_tcPr()
                            shade_obj = OxmlElement("w:shd")
                            shade_obj.set(qn2("w:fill"),"F3FF00")
                            table_cell_properties.append(shade_obj)
                        if row_Cells[4].text == "":
                            cell_xml_element = row_Cells[4]._tc
                            table_cell_properties = cell_xml_element.get_or_add_tcPr()
                            shade_obj = OxmlElement("w:shd")
                            shade_obj.set(qn2("w:fill"),"F3FF00")
                            table_cell_properties.append(shade_obj)
                        cr.execute(f"SELECT * FROM imagesIssues where issuesId='{jx[3]}'")
                        tempValuesImages = cr.fetchone()
                        if tempValuesImages !=None:
                            with open(f"tempimages/{tempValuesImages[1]}{tempValuesImages[2]}","wb") as image:
                                image.write(tempValuesImages[0])
                            paragraph =hdr_Cells[0].paragraphs[0]
                            run = paragraph.runs
                            font = run[0].font
                            font.size= docx.shared.Pt(15)
                            cells = row_Cells[0].paragraphs[0]
                            cells.add_run().add_picture(f"tempimages/{tempValuesImages[1]}{tempValuesImages[2]}",width=docx.shared.Inches(4.5),height=docx.shared.Inches(1))
                            # row_Cells[0].text = str(jx[0])
                        else:
                            cell_xml_element = row_Cells[0]._tc
                            table_cell_properties = cell_xml_element.get_or_add_tcPr()
                            shade_obj = OxmlElement("w:shd")
                            shade_obj.set(qn2("w:fill"),"F3FF00")
                            table_cell_properties.append(shade_obj)

                for row in employees_table.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                
                widths = (docx.shared.Inches(5),docx.shared.Inches(3),docx.shared.Inches(3),docx.shared.Inches(5),docx.shared.Inches(0.5))
                for row in employees_table.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width
                
                for row in employees_table.rows:
                    for cell in row.cells:
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(17)
                if count=="Single":
                    doc.save(f"{file_path}/معلومات المسائلات ل {nameSpicailThingImportant}.docx")
                else:
                    doc.save(f"{file_path}/معلومات المسائلات.docx")

                if what=="Pdf":
                    if count=="Single":
                        # wdFormatPDF = 17
                        in_file = str(os.path.abspath(f"{file_path}/معلومات المسائلات ل {nameSpicailThingImportant}.docx")).replace("c","C")
                        out_file = str(os.path.abspath(f"{file_path}/معلومات المسائلات ل {nameSpicailThingImportant}.pdf")).replace("c","C")
                        
                        # word = comtypes.client.CreateObject('Word.Application')
                        # doc = word.Documents.Open(in_file)
                        # doc.SaveAs(out_file, FileFormat=wdFormatPDF)
                        # doc.Close()
                        # word.Quit()
                        TempdocO = aw.Document(in_file)
                        TempdocO.save(out_file)
                    else:
                        # wdFormatPDF = 17
                        in_file = str(os.path.abspath(f"{file_path}/معلومات المسائلات.docx")).replace("c","C")
                        out_file = str(os.path.abspath(f"{file_path}/معلومات المسائلات.pdf")).replace("c","C") 

                        # word = comtypes.client.CreateObject('Word.Application')
                        # doc = word.Documents.Open(in_file)
                        # doc.SaveAs(out_file, FileFormat=wdFormatPDF)
                        # doc.Close()
                        # word.Quit()
                        TempdocO = aw.Document(in_file)
                        TempdocO.save(out_file)

                    try:
                        if count=="Single":
                            os.remove(f"{file_path}/معلومات المسائلات ل {nameSpicailThingImportant}.docx")
                        else:
                            os.remove(f"{file_path}/معلومات المسائلات.docx")
                    except:
                        pass

                for i in os.listdir("tempImages"):
                    try:
                        os.remove(f"tempImages/{i}")
                    except:
                        pass
                d = QMessageBox(parent=self,text="تم الحفظ بنجاح")
                d.setWindowTitle("نجاح")
                d.setIcon(QMessageBox.Icon.Information)
                d.exec()
        def getExample(self):
            filePath = QFileDialog.getExistingDirectory(self,"Select a Directory")
            if len(filePath)> 0:
                try:
                    os.remove(f"{filePath}/نموذج مسائلة.pdf")
                except:
                    pass
                shutil.copy2("assests/FileExamplePrograme.pdf",f"{filePath}")
                os.rename(f"{filePath}/FileExamplePrograme.pdf",f"{filePath}/نموذج مسائلة.pdf")
                d = QMessageBox(parent=self,text="تم الحفظ بنجاح")
                d.setWindowTitle("نجاح")
                d.setIcon(QMessageBox.Icon.Information)
                d.exec()
        def exportSalaryReport(self):
            for i in self.frame.children():
                i.deleteLater()
            self.frame.setGeometry(self.frame.x(),190,291,200)

            extensiosLabel = QLabel("اختر الصيغة",self.frame)
            extensiosLabel.setStyleSheet("background-color:#389D63;width:250px;color:white")
            extensiosLabel.setFont(QFont("Arial",18))
            extensiosLabel.setAlignment(Qt.AlignmentFlag.AlignCenter)
            extensiosLabel.setGeometry(10,10,270,30)

            self.exportExten = QComboBox(self.frame)
            self.exportExten.setFont(QFont("Arial",15))
            self.exportExten.setGeometry(10,45,270,30)
            self.exportExten.setStyleSheet("background-color:white")

            self.exportExten.addItems(["Word","Pdf"])


            completeExportSalary = QPushButton("تصدير",self.frame)
            completeExportSalary.setFont(QFont("Arial",20))
            completeExportSalary.setGeometry(10,80,270,40)
            completeExportSalary.setStyleSheet("background-color:#389D63;color:white")
            completeExportSalary.clicked.connect(self.writeSalaryWord)

            backButton = QPushButton("العودة",self.frame)
            backButton.setFont(QFont("Arial",20))
            backButton.setGeometry(10,125,270,40)
            backButton.setStyleSheet("background-color:#389D63;color:white")
            backButton.clicked.connect(self.showComponents)

            extensiosLabel.show()
            self.exportExten.show()
            completeExportSalary.show()
            backButton.show()
            self.frame.show()
        def writeSalaryWord(self):
            file_path = QFileDialog.getExistingDirectory(self,"Select a Directory")
            if len(file_path) > 0:
                doc = docx.Document()
                sections = doc.sections
                sections.page_height = 11.69
                sections.page_width = 8.27
                sections = sections[-1]
                sections.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE



                new_width,new_height = sections.page_height,sections.page_width
                sections.page_width = new_width
                sections.page_height = new_height

                sections = doc.sections

                for section in sections:
                    section.top_margin = docx.shared.Cm(0.3)
                    section.bottom_margin = docx.shared.Cm(0.3)
                    section.left_margin = docx.shared.Cm(0.3)
                    section.right_margin = docx.shared.Cm(0.3)
                

                employees_table = doc.add_table(rows=1,cols=4)
                employees_table.style = "Table Grid"
                hdr_Cells = employees_table.rows[0].cells
                hdr_Cells[3].text = "م"
                hdr_Cells[2].text = "اسم المتطوع"
                hdr_Cells[1].text = "السجل المدني"
                hdr_Cells[0].text = "الراتب"
                cr.execute("SELECT name,identy,salary FROM info")                
                b = 0
                tempPictures = []
                cr.execute("SELECT name,identy,salary FROM info")                
                for jx in cr.fetchall():
                    b+=1   
                    row_Cells = employees_table.add_row().cells
                    row_Cells[0].size = docx.shared.Pt(15)
                    row_Cells[1].size = docx.shared.Pt(15)
                    row_Cells[2].size = docx.shared.Pt(15)
                    row_Cells[3].size = docx.shared.Pt(15)

                    row_Cells[3].text = str(b)
                    row_Cells[2].text = str(jx[0])
                    row_Cells[1].text = str(jx[1])
                    row_Cells[0].text = str(jx[2])

                

                for row in employees_table.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

                widths = (docx.shared.Inches(4),docx.shared.Inches(4),docx.shared.Inches(4),docx.shared.Inches(0.5))
                for row in employees_table.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width

                for row in employees_table.rows:
                    for cell in row.cells:
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(17)
                totalSalaryM = 0
                cr.execute("SELECT salary FROM info")
                for i in cr.fetchall():
                    for j in i:
                        totalSalaryM+=int(j)

                doc.add_paragraph()

                employees_table = doc.add_table(rows=1,cols=2)
                employees_table.style = "Table Grid"
                hdr_Cells = employees_table.rows[0].cells
                hdr_Cells[1].text = "اجمالي الرواتب شهريا"
                hdr_Cells[0].text = "ريال {:,.0f}".format((totalSalaryM))

                for row in employees_table.rows:
                    for cell in row.cells:
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(17)

                for row in employees_table.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

                widths = (docx.shared.Inches(6.25),docx.shared.Inches(6.25))
                for row in employees_table.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width

                employees_table = doc.add_table(rows=1,cols=2)
                employees_table.style = "Table Grid"
                hdr_Cells = employees_table.rows[0].cells
                hdr_Cells[1].text = "اجمالي الرواتب سنويا"
                hdr_Cells[0].text = "ريال {:,.0f}".format((totalSalaryM*12))

                for row in employees_table.rows:
                    for cell in row.cells:
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(17)



                for row in employees_table.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

                widths = (docx.shared.Inches(6.25),docx.shared.Inches(6.25))
                for row in employees_table.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width
                doc.save(f"{file_path}//كشف رواتب المتطوعين.docx")
                if self.exportExten.currentText()=="Pdf":
                    # wdFormatPDF = 17
                    
                    in_file = str(os.path.abspath(f"{file_path}/كشف رواتب المتطوعين.docx")).replace("c","C")

                    out_file = str(os.path.abspath(f"{file_path}/كشف رواتب المتطوعين.pdf")).replace("c","C") 

                    TempdocO = aw.Document(in_file)
                    TempdocO.save(out_file)

                    # word = comtypes.client.CreateObject('Word.Application')
                    # doc = word.Documents.Open(in_file)
                    # doc.SaveAs(out_file, FileFormat=wdFormatPDF)
                    # doc.Close()
                    # word.Quit()
                    try:
                        os.remove(f"{file_path}/كشف رواتب المتطوعين.docx")
                    except:
                        pass
                d = QMessageBox(parent=self,text="تم الحفظ بنجاح")
                d.setWindowTitle("نجاح")
                d.setIcon(QMessageBox.Icon.Information)
                d.exec()
        def sendMessages(self):
            self.sendMessagesWindow = tempwindows()
            self.sendMessagesWindow.setFixedSize(400,600)

            self.sendTable = QTableWidget(self.sendMessagesWindow)
            self.sendTable.setColumnCount(3)
            self.sendTable.setHorizontalHeaderLabels(["","الاسم","البريد الالكتروني"])
            self.sendTable.setColumnWidth(0,20)
            self.sendTable.setColumnWidth(1,150)
            self.sendTable.setColumnWidth(2,150)

            self.sendTable.setGeometry(20,20,360,200)
            self.sendTable.setRowCount(0)

            cr.execute("SELECT name,email FROM info")
            for row,i in enumerate(cr.fetchall()):

                item = QTableWidgetItem()
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable|Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(Qt.CheckState.Checked)

                self.sendTable.insertRow(self.sendTable.rowCount())
                self.sendTable.setItem(row,1,QTableWidgetItem(i[0]))
                self.sendTable.setItem(row,2,QTableWidgetItem(i[1]))
                self.sendTable.setItem(row,0,item)

            for i in range(self.sendTable.rowCount()):
                for j in range(self.sendTable.columnCount()):
                    if j!=0:
                        self.sendTable.item(i,j).setFlags(Qt.ItemFlag.ItemIsEditable)


            Label = QLabel("يرجى كتابة الرسالة",self.sendMessagesWindow)
            Label.setFont(QFont("Arial",18))
            Label.move(120,220)

            self.sendMessage = QTextEdit(self.sendMessagesWindow)
            self.sendMessage.setFont(QFont("Arial",16))
            self.sendMessage.setGeometry(20,250,360,200)

            sendButton = QPushButton("ارسال",self.sendMessagesWindow,clicked=self.sendMessagesComplete)
            sendButton.setStyleSheet("background-color:#44C17A;width:250px;color:black")
            sendButton.setFont(QFont("Arial",20))
            sendButton.move(70,470)

            self.sendMessagesWindow.show()
        def sendMessagesComplete(self):
            self.emails = []
            self.names = []
            self.bodyMessage = self.sendMessage.toPlainText()
            for i in range(self.sendTable.rowCount()):
                if self.sendTable.item(i,0).checkState()==Qt.CheckState.Checked:
                    self.emails.append(self.sendTable.item(i,2).text())
                    self.names.append(self.sendTable.item(i,1).text())
                    
            d = QMessageBox(parent=self,text=f"اجمالي عدد الرسائل {len(self.emails)} هل انت متأكد من الارسال")
            d.setWindowTitle(title)
            d.setIcon(QMessageBox.Icon.Information)
            d.setStandardButtons(QMessageBox.StandardButton.Ok|QMessageBox.StandardButton.Cancel)
            ret = d.exec()
            if ret ==QMessageBox.StandardButton.Ok:
                try:
                    self.sendMessagesWindow.destroy()
                except:
                    pass
                self.loadingScreen = tempwindows()
                self.loadingScreen.setFixedSize(300,300)

                sendLabel = QLabel("جاري الارسال ... يرجى الانتظار",self.loadingScreen)
                sendLabel.setStyleSheet("color:green")
                sendLabel.setFont(QFont("Arial",15))
                sendLabel.move(40,20)

                self.labelMovie = QLabel(self.loadingScreen)
                self.movie = QMovie("assests/loading.gif")
                self.movie.setScaledSize(QSize(100,100))
                self.labelMovie.setMovie(self.movie)
                self.labelMovie.move(95,100)

                self.startAnimation()
                timer = QTimer(self.loadingScreen)
                timer.singleShot(3000,self.sendMessagesAction)
                self.loadingScreen.show()
        def startAnimation(self):
            self.movie.start()
        def stopAnimation(self):
            self.movie.stop()
            self.loadingScreen.destroy()
        def sendMessagesAction(self):
            total = 0
            succses = 0
            fail = 0

            email_sender = "abodi3313@gmail.com"
            email_pass = "dduviygcxsmyxekr"
            subjet = "تعميم"

            em = EmailMessage()
            em['From']=email_sender
            em['subject'] = subjet
            em.set_content(self.bodyMessage)
            context = ssl.create_default_context()
            doc = docx.Document()

            sections = doc.sections
            sections.page_height = 11.69
            sections.page_width = 8.27
            sections = sections[-1]
            sections.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE

            new_width,new_height = sections.page_height,sections.page_width
            sections.page_width = new_width
            sections.page_height = new_height

            sections = doc.sections

            for section in sections:
                section.top_margin = docx.shared.Cm(0.3)
                section.bottom_margin = docx.shared.Cm(0.3)
                section.left_margin = docx.shared.Cm(0.3)
                section.right_margin = docx.shared.Cm(0.3)

            
            emails_table = doc.add_table(rows=1,cols=4)
            emails_table.style = "Table Grid"
            hdr_Cells = emails_table.rows[0].cells
            hdr_Cells[3].text = "م"
            hdr_Cells[2].text = "اسم المتطوع"
            hdr_Cells[1].text = "الايميل"
            hdr_Cells[0].text = "حالة الارسال"


            for n,email in enumerate(self.emails):
                row_Cells = emails_table.add_row().cells
                row_Cells[3].text = str(n+1)
                row_Cells[2].text = str(self.names[n])
                row_Cells[1].text = str(email)

                try:
                    del em['To']
                    em['To']=email
                    with smtplib.SMTP_SSL('smtp.gmail.com',465,context=context) as smtp:
                        smtp.login(email_sender, email_pass)
                        smtp.sendmail(email_sender, email, em.as_string())
                    row_Cells[0].text = "نجح الارسال"
                    succses +=1
                except:
                    row_Cells[0].text = "فشل الارسال"
                    fail +=1
                total+=1
            for row in emails_table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

            widths = (docx.shared.Inches(4),docx.shared.Inches(4),docx.shared.Inches(4),docx.shared.Inches(0.5))

            for row in emails_table.rows:
                for idx, width in enumerate(widths):
                    row.cells[idx].width = width
            
            for row in emails_table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= docx.shared.Pt(17)
            
            doc.add_paragraph()

            emails_table = doc.add_table(rows=1,cols=3)
            emails_table.style = "Table Grid"
            hdr_Cells = emails_table.rows[0].cells

            hdr_Cells[2].text = "اجمالي عدد الرسائل"
            hdr_Cells[1].text = "الرسائل الناجحة"
            hdr_Cells[0].text = "الرسائل الفاشلة"

            row_Cells = emails_table.add_row().cells
            row_Cells[2].text = str(total)
            row_Cells[1].text = str(succses)
            row_Cells[0].text = str(fail)

            for row in emails_table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

            widths = (docx.shared.Inches(4.25),docx.shared.Inches(4.25),docx.shared.Inches(4))

            for row in emails_table.rows:
                for idx, width in enumerate(widths):
                    row.cells[idx].width = width
            
            for row in emails_table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= docx.shared.Pt(17)

            self.stopAnimation()
            d = QMessageBox(parent=self,text="تم الارسال بنجاح هل تريد حفظ تقرير الارسال")
            d.setWindowTitle("نجاح")
            d.setIcon(QMessageBox.Icon.Information)
            d.setStandardButtons(QMessageBox.StandardButton.Ok|QMessageBox.StandardButton.Cancel)
            ret = d.exec()
            if ret==QMessageBox.StandardButton.Ok:
                file_path = QFileDialog.getExistingDirectory(self,"Select a Directory")
                if len(file_path) > 0:
                    doc.save(f"{file_path}/تقرير الارسال.docx")

                    in_file = str(os.path.abspath(f"{file_path}/تقرير الارسال.docx")).replace("c","C")
                    out_file = str(os.path.abspath(f"{file_path}/تقرير الارسال.pdf")).replace("c","C")
                    TempdocO = aw.Document(in_file)
                    TempdocO.save(out_file)

                    os.remove(f"{file_path}/تقرير الارسال.docx")
                    d = QMessageBox(parent=self,text="تم الحفظ بنجاح")
                    d.setWindowTitle("نجاح")
                    d.setIcon(QMessageBox.Icon.Information)
                    d.exec()
    if __name__ == "__main__":
        mainToRunApp()
else:
    ctypes.windll.shell32.ShellExecuteW(None, "runas", sys.executable, " ".join(sys.argv), None, 1)


